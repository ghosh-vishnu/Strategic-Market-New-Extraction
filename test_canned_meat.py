import sys
import os

# Add backend directory to path
backend_path = os.path.join(os.path.dirname(__file__), 'backend')
sys.path.insert(0, backend_path)
sys.path.insert(0, os.path.join(backend_path, 'converter', 'utils'))

from converter.utils.extractor import extract_title
from docx import Document
from converter.utils.extractor import remove_emojis
import re

# Test file path
docx_path = r'C:\Users\Vishnu\Desktop\Project\Strategic-Market\Canned Meat Market.docx'

print("=" * 80)
print("Testing extract_title for: Canned Meat Market.docx")
print("=" * 80)

# First, analyze the document structure
doc = Document(docx_path)
filename = os.path.splitext(os.path.basename(docx_path))[0]

print("\n📄 First 40 Paragraphs:")
print("-" * 80)
for i, para in enumerate(doc.paragraphs[:40]):
    text = para.text.strip()
    if text:
        clean_text = remove_emojis(text)
        print(f"\n[{i}] {clean_text[:200]}")
        
        # Highlight potential title patterns
        text_lower = clean_text.lower()
        if 'canned meat' in text_lower:
            print("   ⭐ Contains market name keywords")
        if 'report title' in text_lower or 'full title' in text_lower:
            print("   🔍 REPORT TITLE HEADER")
        if 'global' in text_lower and 'market' in text_lower:
            print("   🌍 GLOBAL MARKET pattern")
        if re.search(r'by\s+(application|type|product|end[-\s]*user|region|geography)', text_lower):
            print("   📊 SEGMENTATION pattern found")
        if 'forecast' in text_lower and re.search(r'20\d{2}', clean_text):
            print("   📅 FORECAST with year")

# Check tables
print("\n\n📊 Tables Analysis:")
print("-" * 80)
for table_idx, table in enumerate(doc.tables[:3]):
    print(f"\nTable {table_idx}:")
    for row_idx, row in enumerate(table.rows[:5]):
        row_text = []
        for cell in row.cells:
            cell_text = cell.text.strip()[:60]
            row_text.append(cell_text)
        print(f"  Row {row_idx}: {' | '.join(row_text)}")
        if 'report title' in ' '.join(row_text).lower():
            print("     🔍 REPORT TITLE found in table!")

# Test the function
print("\n\n🔧 Testing extract_title function:")
print("-" * 80)
result = extract_title(docx_path)
print(f"Result: {result}")
print(f"Length: {len(result)} chars")
print(f"Word count: {len(result.split())} words")

if result == "Title Not Available":
    print("\n❌ FAILED: Title not extracted!")
else:
    print("\n✅ SUCCESS: Title extracted successfully!")

