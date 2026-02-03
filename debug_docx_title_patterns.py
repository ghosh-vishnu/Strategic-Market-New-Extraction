import argparse
import os
import re

from docx import Document


def norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip())


def main() -> int:
    parser = argparse.ArgumentParser(description="Inspect a .docx for title/segmentation patterns.")
    parser.add_argument("docx_path", help="Path to .docx")
    parser.add_argument("--paras", type=int, default=80, help="How many paragraphs to print (from start).")
    parser.add_argument("--tables", type=int, default=5, help="How many tables to scan (from start).")
    args = parser.parse_args()

    docx_path = os.path.abspath(args.docx_path)
    doc = Document(docx_path)

    print(f"File: {docx_path}")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")

    triggers = [
        ("report title", re.compile(r"(report title|full report title|full title)", re.I)),
        ("global market", re.compile(r"(^|\\b)(the\\s+)?global\\b.*\\bmarket\\b", re.I)),
        ("segmentation", re.compile(r"\\bby\\s+(application|product\\s+type|type|end[-\\s]*user|region|geography|distribution\\s+channel)\\b", re.I)),
        ("forecast", re.compile(r"\\bforecast\\b", re.I)),
        ("segment revenue", re.compile(r"segment\\s+revenue\\s+estimation", re.I)),
        ("year range", re.compile(r"20\\d{2}\\s*[\\-–]\\s*20\\d{2}")),
    ]

    print("\n--- First paragraphs ---")
    shown = 0
    for i, para in enumerate(doc.paragraphs[: args.paras]):
        t = norm(para.text)
        if not t:
            continue
        marks = []
        for name, rx in triggers:
            if rx.search(t):
                marks.append(name)
        mark_txt = f"  [{', '.join(marks)}]" if marks else ""
        print(f"[{i}] {t[:220]}{mark_txt}")
        shown += 1
        if shown >= args.paras:
            break

    print("\n--- Table scan (first matches) ---")
    for t_idx, table in enumerate(doc.tables[: args.tables]):
        for r_idx, row in enumerate(table.rows):
            row_text = " | ".join(norm(c.text)[:80] for c in row.cells)
            if not row_text.strip():
                continue
            low = row_text.lower()
            if any(k in low for k in ["report title", "full title", "global", "market", "by ", "forecast", "segment revenue"]):
                print(f"[table {t_idx} row {r_idx}] {row_text[:300]}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())

