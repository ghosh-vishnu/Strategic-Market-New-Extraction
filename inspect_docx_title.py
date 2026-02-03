"""Inspect where title / Report Title (Long-Form) appears in a docx."""
import os
import sys
import re

root = os.path.dirname(os.path.abspath(__file__))
from docx import Document

def main():
    path_bad = os.path.join(root, "Extracted", "Creatine Kinase Reagent Market.docx")
    path_arb = os.path.join(root, "Extracted", "Arbovirus Testing Market.docx")
    path_ok = os.path.join(root, "Okay", "Blood Collection Tubes Market.docx")
    for label, path in [("Extracted Arbovirus", path_arb), ("Extracted (bad)", path_bad), ("Okay (ref)", path_ok)]:
        if not os.path.isfile(path):
            print(label, path, "not found")
            continue
        doc = Document(path)
        print("\n=== ", label, " ===")
        print("Paragraphs:", len(doc.paragraphs))
        rx = re.compile(r"report title|full title|long.form|segment revenue estimation|market segmentation and forecast|^form\)", re.I)
        for i, p in enumerate(doc.paragraphs):
            t = (p.text or "").strip()
            if not t:
                continue
            if rx.search(t) or (i < 15 and "market" in t.lower()) or (160 <= i <= 175):
                print(i, ":", t[:180].encode("ascii", "replace").decode())
        # Last 30 paras (often A.1 is near end)
        print("--- last 30 non-empty ---")
        count = 0
        for i in range(len(doc.paragraphs) - 1, -1, -1):
            if count >= 30:
                break
            t = (doc.paragraphs[i].text or "").strip()
            if not t:
                continue
            count += 1
            print(i, ":", t[:160].encode("ascii", "replace").decode())

if __name__ == "__main__":
    main()
