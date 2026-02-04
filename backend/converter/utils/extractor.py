from docx import Document
from datetime import date
import json
import html
import re
import os
import pandas as pd
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn
from docx.text.run import Run
import concurrent.futures
import threading
from functools import lru_cache

# ------------------- Helpers -------------------
DASH = "–"  # en-dash for year ranges
EXCEL_CELL_LIMIT = 32767  # Excel max char limit per cell

# ------------------- Performance Optimizations -------------------

# Thread-safe cache for common patterns
_pattern_cache = {}
_cache_lock = threading.Lock()

@lru_cache(maxsize=128)
def _get_cached_pattern(pattern_key: str, pattern: str):
    """Cache compiled regex patterns for better performance."""
    return re.compile(pattern, re.I | re.X)

def _get_pattern(pattern_key: str, pattern: str):
    """Get cached regex pattern or create new one."""
    with _cache_lock:
        if pattern_key not in _pattern_cache:
            _pattern_cache[pattern_key] = re.compile(pattern, re.I | re.X)
        return _pattern_cache[pattern_key]

def remove_emojis(text: str) -> str:
    """Universal emoji remover."""
    emoji_pattern = re.compile(
        "[" 
        "\U0001F600-\U0001F64F"  # emoticons
        "\U0001F300-\U0001F5FF"  # symbols & pictographs
        "\U0001F680-\U0001F6FF"  # transport & map
        "\U0001F700-\U0001F77F"  # alchemical
        "\U0001F780-\U0001F7FF"  # geometric
        "\U0001F800-\U0001F8FF"  # arrows
        "\U0001F900-\U0001F9FF"  # supplemental
        "\U0001FA00-\U0001FAFF"  # chess, symbols
        "\U00002600-\U000026FF"  # misc symbols
        "\U00002700-\U000027BF"  # dingbats
        "\U00002B00-\U00002BFF"  # arrows & symbols
        "\U0001F1E0-\U0001F1FF"  # flags
       "\U00010000-\U0010ffff"
        "]+", flags=re.UNICODE
    )
    return emoji_pattern.sub(r'', text or "")

# ------------------- Normalization ------------------- 
def _norm(s: str) -> str:
    s = remove_emojis(s or "")
    # Normalize common dash/encoding issues in DOCX text.
    # Some generated documents contain U+FFFD (�) where an en-dash (–) should be.
    s = s.replace("\uFFFD", DASH)  # � -> –
    s = s.replace("\u2014", DASH)  # em-dash -> en-dash
    return re.sub(r"\s+", " ", s.strip())

def _inline_title(text: str) -> str:
    m = re.split(r"[:\-–]", text, maxsplit=1)
    if len(m) > 1:
        right = m[1].strip()
        if right and not re.match(r'^\s*(?:[A-Za-z]\.)?(?:\d+(?:\.\d+)*)?[\.\)]?\s*(?:report\s*title|full\s*title|full\s*report\s*title|title\s*\(long[-\s]*form\))[\s:–-]*$', right, re.I):
            # Reject fragments from splitting on hyphen in "Long-Form" (e.g. "Form)")
            if len(right) >= 25 and ("market" in right.lower() or " by " in right.lower() or "segment revenue" in right.lower()):
                return right
    return ""

def _year_range_present(text: str) -> bool:
    return bool(re.search(r"20\d{2}\s*[\-–]\s*20\d{2}", text))


def _clean_final_title(title: str) -> str:
    """Remove common doc artifacts from extracted title."""
    if not title or len(title) < 5:
        return title
    # Remove '(Structured):' or '(Structured): ' prefix
    title = re.sub(r"^\s*\(structured\)\s*:\s*", "", title, flags=re.I).strip()
    # Remove duplicate market name: "Global Pediatric Catheter Market Pediatric Catheter Market By" -> "Global Pediatric Catheter Market By"
    title = re.sub(
        r"\b(Global\s+)?(.+?Market)\s+\2(\s+By\s+)",
        r"\1\2\3",
        title,
        flags=re.I,
    )
    # Remove parentheticals that are narrative (e.g. "This is the most critical dimension of segmentation", "In 2024, SPECT/...")
    title = re.sub(r"\s*\([^)]*this is the[^)]*\)", "", title, flags=re.I)
    title = re.sub(r"\s*\(in 2024[^)]*\)", "", title, flags=re.I)
    title = re.sub(r"\s*\([^)]*dimension of segmentation[^)]*\)", "", title, flags=re.I)
    title = re.sub(r"\s{2,}", " ", title).strip()
    return title


def _ensure_filename_start_and_year(title: str, filename: str) -> str:
    # Normalize for comparison
    title_lower = title.lower()
    filename_lower = filename.lower()
    filename_normalized = filename_lower.replace('-', ' ').replace('_', ' ')
    title_normalized = title_lower.replace('-', ' ').replace('_', ' ')
    
    # Extract key words from filename (excluding common words)
    filename_keywords = [w for w in filename_normalized.split() if w not in ['market', 'the', 'and', 'or']]
    
    # Check if title already contains filename keywords
    matching_keywords = sum(1 for kw in filename_keywords if kw in title_normalized)
    
    # Don't prepend when title looks complete: has "Market", " By ", and (2+ filename keywords or 2+ "By" segments)
    by_count = len(re.findall(r"\bby\s+\w", title_lower))
    looks_complete = (
        " by " in title_lower
        and "market" in title_lower
        and (matching_keywords >= min(2, len(filename_keywords)) or by_count >= 2)
    )
    # Only prepend filename if title doesn't already contain it (and isn't a complete doc title)
    if not looks_complete and matching_keywords < min(3, len(filename_keywords)) and not title_lower.startswith(filename_lower):
        title = f"{filename} {title}"
    
    if not _year_range_present(title):
        title = f"{title} 2024{DASH}2030"
    return _clean_final_title(_norm(title))

# ✅ Detect list items
def is_list_item(para):
    pPr = para._p.pPr
    if pPr is not None and pPr.numPr is not None:
        return True
    return False

def get_list_style_type(para, text):
    """Detect the list style type (bullet, number, circle, etc.)"""
    # Check for bullet characters
    if '•' in text or '*' in text:
        return "bullet"
    elif '○' in text or '◦' in text:
        return "circle"
    elif '–' in text or '—' in text:
        return "dash"
    elif re.match(r'^\d+[\.\)]', text):
        return "number"
    elif '▪' in text or '▫' in text:
        return "square"
    else:
        # Default to bullet if no specific character found
        return "bullet"

def is_main_level_item(para, text, parent_list_style):
    """Determine if an item should be at main level based on style comparison"""
    current_style = get_list_style_type(para, text)
    
    # If no parent style, it's main level
    if parent_list_style is None:
        return True
    
    # If styles are different, it should be main level
    if current_style != parent_list_style:
        return True
    
    # If it's a bold item that doesn't end with colon, it's likely main level
    if '<b>' in text or '<strong>' in text:
        if not text.strip().endswith(':'):
            return True
    
    # If it's a known main level item pattern
    main_level_patterns = [
        "Benchmarking of Market Leaders",
        "Recent Product Developments and Approvals", 
        "Strategic Collaborations and M&A",
        "Market Share Analysis",
        "Competitive Strategies"
    ]
    
    for pattern in main_level_patterns:
        if pattern in text:
            return True
    
    return False

# ------------------- Convert Paragraph to HTML -------------------
def runs_to_html(runs):
    """Convert Word runs (bold/italic) to inline HTML with hyperlink support."""
    parts = []
    for run in runs:
        txt = remove_emojis(run.text.strip())
        if not txt:
            continue

        # hyperlink detection
        if run._element.xpath("ancestor::w:hyperlink"):
            rId = run._element.xpath("ancestor::w:hyperlink/@r:id")
            if rId:
                try:
                    link = run.part.rels[rId[0]].target_ref
                    parts.append(f'<a href="{link}">{txt}</a>')
                except Exception:
                    parts.append(txt)
            else:
                parts.append(txt)
        elif run.bold and run.italic:
            parts.append(f"<b><i>{txt}</i></b>")
        elif run.bold:
            parts.append(f"<b>{txt}</b>")
        elif run.italic:
            parts.append(f"<i>{txt}</i>")
        else:
            parts.append(txt)
    return " ".join(parts).strip()

def extract_table_with_style(table):
    """Extract table with exact report styling using inline CSS"""
    html_parts = []
    # Exact report table styling with inline CSS
    html_parts.append('<table cellspacing=0 style="border-collapse:collapse; width:100%">')
    html_parts.append('<tbody>')
    
    for row_idx, row in enumerate(table.rows):
        html_parts.append('<tr>')
        
        for cell_idx, cell in enumerate(row.cells):
            cell_text = " ".join(
                runs_to_html(para.runs) for para in cell.paragraphs
            ).strip()
            
            # Determine if this is header row
            is_header = row_idx == 0
            
            # Determine if this row should have alternating background
            is_alternating_row = row_idx % 2 == 1  # Odd rows (1, 3, 5, etc.)
            
            # Determine if this is the last cell in the row
            is_last_cell = cell_idx == len(row.cells) - 1
            
            # Build cell style based on position and row type
            cell_style_parts = []
            
            # Background color
            if is_header:
                cell_style_parts.append('background-color:#4472c4')
            elif is_alternating_row:
                cell_style_parts.append('background-color:#d9e2f3')
            else:
                cell_style_parts.append('background-color:#ffffff')
            
            # Border styling
            border_color = '#4472c4' if is_header else '#8eaadb'
            
            # Top border
            if row_idx == 0:
                cell_style_parts.append('border-top:1px solid #4472c4')
            else:
                cell_style_parts.append('border-top:none')
            
            # Bottom border
            cell_style_parts.append(f'border-bottom:1px solid {border_color}')
            
            # Left border
            cell_style_parts.append(f'border-left:1px solid {border_color}')
            
            # Right border
            if is_last_cell:
                cell_style_parts.append(f'border-right:1px solid {border_color}')
            else:
                cell_style_parts.append('border-right:none')
            
            # Other styling
            cell_style_parts.extend([
                'vertical-align:top',
                'width:195px' if cell_idx == 0 else 'width:370px'
            ])
            
            cell_style = '; '.join(cell_style_parts)
            
            # Wrap content in paragraph with strong tags
            content = f'<p><strong>{cell_text}</strong></p>'
            
            html_parts.append(f'<td style="{cell_style}">{content}</td>')
        
        html_parts.append('</tr>')
    
    html_parts.append('</tbody>')
    html_parts.append('</table>')
    return '\n'.join(html_parts)



# ------------------- TOC Extraction -------------------
def determine_toc_logic(doc):
    """Determine which logic to use based on document structure"""
    executive_summary_bold = False
    executive_summary_in_list = False
    first_line_after_bold = False
    first_line_after_in_list = False
    has_nested_lists = False
    para_count = 0
    found_executive_summary = False
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Check if this is Executive Summary
        if "Executive Summary" in text:
            found_executive_summary = True
            # Check for bold (including <strong> tags)
            executive_summary_bold = any(run.bold for run in para.runs if run.text.strip()) or '<strong>' in text or '<b>' in text
            executive_summary_in_list = is_list_item(para) or any(char in text for char in ['•', '-', '–', '○', '◦', '‣', '▪', '▫', '*', '+']) or re.match(r'^\d+[\.\)]', text)
            print(f"DEBUG: Executive Summary found: '{text[:50]}...' - Bold: {executive_summary_bold}, In List: {executive_summary_in_list}")
            continue
            
        # Only check lines after Executive Summary
        if not found_executive_summary:
            continue
            
        para_count += 1
        # Check for bold (including <strong> tags)
        is_bold = any(run.bold for run in para.runs if run.text.strip()) or '<strong>' in text or '<b>' in text
        # Check for list items (including Word's list formatting)
        is_in_list = (
            is_list_item(para) or  # Check Word's list formatting first
            any(char in text for char in ['•', '-', '–', '○', '◦', '‣', '▪', '▫', '*', '+']) or 
            re.match(r'^\d+[\.\)]', text)
        )
        
        # Check first line after Executive Summary
        if para_count == 1:
            first_line_after_bold = is_bold
            first_line_after_in_list = is_in_list
            print(f"DEBUG: First line after Executive Summary: '{text[:50]}...' - Bold: {is_bold}, In List: {is_in_list}")
        
        # Check for nested lists (parent/child structure)
        # LOGIC 2 is detected when we have a mix of main headings and list items
        # For now, we'll disable automatic nested list detection and rely on the user's input
        # to determine which logic to use
    
    # Logic detection
    print(f"DEBUG: Detection results - para_count: {para_count}, has_nested_lists: {has_nested_lists}")
    print(f"DEBUG: executive_summary_bold: {executive_summary_bold}, executive_summary_in_list: {executive_summary_in_list}")
    print(f"DEBUG: first_line_after_bold: {first_line_after_bold}, first_line_after_in_list: {first_line_after_in_list}")
    
    if para_count >= 1:
        # LOGIC 2: Parent-child structure (Executive Summary bold+list, first line after non-bold+list)
        if executive_summary_bold and executive_summary_in_list and first_line_after_in_list and not first_line_after_bold:
            print("DEBUG: Selected LOGIC 2 (Parent-child structure: Executive Summary bold+list, first line after non-bold+list)")
            return 2
        
        # LOGIC 1: Executive Summary bold, first line after in list (non-bold) but not parent-child
        elif executive_summary_bold and first_line_after_in_list and not first_line_after_bold:
            print("DEBUG: Selected LOGIC 1 (Executive Summary bold, first line after list non-bold)")
            return 1
        
        # LOGIC 3: Executive Summary bold, first line after in list (bold)
        elif executive_summary_bold and first_line_after_in_list and first_line_after_bold:
            print("DEBUG: Selected LOGIC 3 (Executive Summary bold, first line after list bold)")
            return 3
        
        # Special case: If Executive Summary is bold and first line after is bold but not detected as list,
        # but we can see from context that it should be LOGIC 3 (based on user's image)
        elif executive_summary_bold and first_line_after_bold and not first_line_after_in_list:
            print("DEBUG: Selected LOGIC 3 (Executive Summary bold, first line after bold but not detected as list - forcing LOGIC 3)")
            return 3
    
    # Default to logic 1
    print("DEBUG: Selected LOGIC 1 (default)")
    return 1

def extract_toc(docx_path):
    doc = Document(docx_path)
    html_output = []
    capture = False
    inside_list = False
    list_depth = 0  # Track nesting depth
    previous_was_bold = False
    in_nested_context = False
    parent_list_style = None  # Track parent list style
    current_list_style = None  # Track current list style
    
    # Determine which logic to use for this file
    logic_type = determine_toc_logic(doc)
    print(f"DEBUG: Logic type determined: {logic_type}")  # Debug output

    def clean_heading(text):
        """Clean heading text by removing numbering, bullets, and extra spaces"""
        text = remove_emojis(text.strip())
        # Remove numbering patterns like "1.", "1.1", "1.1.1", etc.
        text = re.sub(r'^\d+(\.\d+)*[\.\)]\s*', '', text)
        # Remove bullet points
        text = re.sub(r'^[•\-–]\s*', '', text)
        # Remove extra spaces
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    def is_bold_text(para):
        """Check if paragraph has bold text (including <strong> tags)"""
        if para.runs:
            # Check for run.bold
            has_bold_run = any(run.bold for run in para.runs if run.text.strip())
            if has_bold_run:
                return True
            
            # Also check if the text contains <strong> tags
            text = para.text.strip()
            if '<strong>' in text or '<b>' in text:
                return True
                
        return False

    def is_heading(para):
        """Check if paragraph is a heading based on style or pattern"""
        style_name = getattr(para.style, "name", "").lower()
        if "heading" in style_name:
            return True
        # Check for numbered patterns like "1. Title", "1.1 Subtitle"
        if re.match(r'^\d+(\.\d+)*[\.\)]\s+', para.text.strip()):
            return True
        return False

    def is_subheading(para):
        """Check if paragraph is a subheading (level 2 or deeper)"""
        style_name = getattr(para.style, "name", "").lower()
        if "heading" in style_name:
            level = para.style.name.replace("Heading", "").strip()
            if level.isdigit() and int(level) >= 3:
                return True
        # Check for deeper numbering patterns like "1.1", "1.1.1", etc.
        if re.match(r'^\d+\.\d+', para.text.strip()):
            return True
        return False

    def runs_to_html_with_links(runs):
        """Convert Word runs to HTML with proper formatting and links"""
        parts = []
        for run in runs:
            txt = remove_emojis(run.text.strip())
            if not txt:
                continue

            # Check for hyperlinks
            if run._element.xpath("ancestor::w:hyperlink"):
                rId = run._element.xpath("ancestor::w:hyperlink/@r:id")
                if rId:
                    try:
                        link = run.part.rels[rId[0]].target_ref
                        parts.append(f'<a href="{link}">{txt}</a>')
                    except Exception:
                        parts.append(txt)
                else:
                    parts.append(txt)
            elif run.bold and run.italic:
                parts.append(f"<b><i>{txt}</i></b>")
            elif run.bold:
                parts.append(f"<b>{txt}</b>")
            elif run.italic:
                parts.append(f"<i>{txt}</i>")
            else:
                parts.append(txt)
        return " ".join(parts).strip()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        cleaned_text = clean_heading(text)
        low = cleaned_text.lower()
        is_bold = is_bold_text(para)

        # Start condition: Look for "Executive Summary" (ignore numbering/bullets)
        if not capture and "Executive Summary" in text:
            capture = True
            # Process Executive Summary itself
            if is_bold:
                heading_text = clean_heading(text)
                # If Executive Summary line contains inline bullets (•), split into heading + list
                if '•' in text:
                    parts = [p.strip() for p in text.split('•') if p.strip()]
                    if parts:
                        # First part as heading (cleaned)
                        html_output.append(f"\n<strong>{clean_heading(parts[0])}</strong>")
                        # Remaining parts as bullets
                        bullets = parts[1:]
                        if bullets:
                            html_output.append("<ul>")
                            for item in bullets:
                                item_clean = remove_emojis(item).strip()
                                if item_clean:
                                    html_output.append(f"<li><p>{item_clean}</p></li>")
                            html_output.append("</ul>")
                else:
                    if heading_text:
                        # For headings, keep only <strong> tags, remove <b> tags
                        heading_text = heading_text.replace('<b>', '').replace('</b>', '')
                        html_output.append(f"\n<strong>{heading_text}</strong>")
            continue

        # Only process content after Executive Summary is found
        if capture:
            # Apply single logic based on document structure
            if logic_type == 1:
                # LOGIC 1: Bold = H2, Non-bold = p tags (preserve nested structure)
                print(f"DEBUG LOGIC 1: Processing '{text[:30]}...' - is_bold: {is_bold}")
                
                # Check if this is a nested list item - rely primarily on Word's list formatting
                is_word_list_item = is_list_item(para)
                # Also check for common list patterns as fallback
                has_bullet_chars = any(char in text for char in ['•', '–', '○', '◦', '‣', '▪', '▫', '*', '+'])
                has_numbering = re.match(r'^\d+[\.\)]', text)
                is_nested_list = is_word_list_item or has_bullet_chars or has_numbering
                
                if is_bold:
                    # Check if this bold text is in a list (visible indicators OR Word's internal formatting)
                    is_in_list = has_bullet_chars or has_numbering or is_word_list_item
                    
                    # Debug: Check what's being detected
                    print(f"DEBUG: Bold text '{text[:50]}...' - has_bullet_chars: {has_bullet_chars}, has_numbering: {has_numbering}, is_word_list_item: {is_word_list_item}, is_in_list: {is_in_list}")
                    
                    if is_in_list:
                        # This is bold text within a list item - keep it as part of the list
                        formatted_content = runs_to_html_with_links(para.runs)
                        
                        # Check if this should be a parent item (ends with colon)
                        is_parent_item = (
                            ":" in formatted_content and formatted_content.strip().endswith(":")
                        ) and "List of Figures" not in "\n".join(html_output[-10:])
                        
                        if is_parent_item:
                            # This is a parent item that should have nested children
                            if inside_list:
                                # Close any existing nested structure
                                for _ in range(list_depth):
                                    html_output.append("</ul>")
                                list_depth = 0
                                inside_list = False
                            
                            html_output.append("<ul>")
                            html_output.append(f"<li><p><strong>{formatted_content}</strong></p>")
                            html_output.append("<ul>")  # Start nested list for children
                            inside_list = True
                            list_depth = 2  # We have main list + nested list
                            print(f"DEBUG LOGIC 1: Added bold parent item with nested list: {formatted_content[:30]}...")
                        else:
                            # This is a bold list item (not a parent)
                            if not inside_list:
                                html_output.append("<ul>")
                                inside_list = True
                                list_depth = 1
                            
                            html_output.append(f"<li><p>{formatted_content}</p></li>")
                            print(f"DEBUG LOGIC 1: Added bold list item: {formatted_content[:30]}...")
                    else:
                        # Bold text is NOT in list - treat as heading
                        # Close any open lists first
                        if inside_list:
                            # Close all open lists based on depth
                            for _ in range(list_depth):
                                html_output.append("</ul>")
                            inside_list = False
                            list_depth = 0
                        
                        heading_text = clean_heading(text)
                        if heading_text:
                            # For headings, keep only <strong> tags, remove <b> tags
                            heading_text = heading_text.replace('<b>', '').replace('</b>', '')
                            html_output.append(f"\n<strong>{heading_text}</strong>")
                            print(f"DEBUG LOGIC 1: Added <strong> for bold heading (not in list): {heading_text[:30]}...")
                else:
                    # Non-bold text - check if it's a nested list item
                    if is_nested_list:
                        # This is a nested list item - check list style for nesting logic
                        formatted_content = runs_to_html_with_links(para.runs)
                        
                        # Get current list style (bullet type)
                        current_list_style = "bullet"  # Default
                        if has_bullet_chars:
                            if '•' in text or '*' in text:
                                current_list_style = "bullet"
                            elif '○' in text or '◦' in text:
                                current_list_style = "circle"
                            elif '–' in text or '–' in text:
                                current_list_style = "dash"
                        elif has_numbering:
                            current_list_style = "number"
                        
                        # Check if this should be a parent item (ends with colon)
                        is_parent_item = (
                            ":" in formatted_content and formatted_content.strip().endswith(":")
                        ) and "List of Figures" not in "\n".join(html_output[-10:])
                        
                        if is_parent_item:
                            # This is a parent item that should have nested children
                            if inside_list:
                                # Close any existing nested structure
                                for _ in range(list_depth):
                                    html_output.append("</ul>")
                                list_depth = 0
                                inside_list = False
                            
                            html_output.append("<ul>")
                            html_output.append(f"<li><p><strong>{formatted_content}</strong></p>")
                            html_output.append("<ul>")  # Start nested list for children
                            inside_list = True
                            list_depth = 2  # We have main list + nested list
                            parent_list_style = current_list_style  # Set parent style
                            print(f"DEBUG LOGIC 1: Added parent item with nested list: {formatted_content[:30]}...")
                        else:
                            # Check if this should be a new main section (like "Strategy Analysis")
                            is_new_main_section = ("Strategy Analysis:" in formatted_content )
                            
                            if is_new_main_section:
                                # Close any existing nested structure first
                                if inside_list and list_depth > 1:
                                    html_output.append("</ul>")  # Close nested list
                                    html_output.append("</li>")  # Close parent item
                                    html_output.append("</ul>")  # Close main list
                                    inside_list = False
                                    list_depth = 0
                                
                                # Add as new main list item
                                if not inside_list:
                                    html_output.append("<ul>")
                                    inside_list = True
                                    list_depth = 1
                                
                                html_output.append(f"<li><p>{formatted_content}</p></li>")
                                print(f"DEBUG LOGIC 1: Added main list item: {formatted_content[:30]}...")
                            else:
                                # Check if this should be a main level item
                                should_be_main_level = is_main_level_item(para, formatted_content, parent_list_style)
                                
                                if should_be_main_level and inside_list and list_depth > 1:
                                    # This should be at main level - close child list first
                                    html_output.append("</ul>")  # Close child list
                                    html_output.append("</li>")  # Close parent item
                                    list_depth = 1  # Back to main list level
                                    parent_list_style = current_list_style  # Update parent style
                                    print(f"DEBUG LOGIC 1: Main level item detected - closing child list: {formatted_content[:30]}...")
                                
                                # Add as main list item
                                if not inside_list:
                                    html_output.append("<ul>")
                                    inside_list = True
                                    list_depth = 1
                                    parent_list_style = current_list_style  # Set parent style
                                
                                html_output.append(f"<li><p>{formatted_content}</p></li>")
                                print(f"DEBUG LOGIC 1: Added list item: {formatted_content[:30]}...")
                    elif ":" in formatted_content and formatted_content.strip().endswith(":"):
                        # This is a parent item that should have nested children (not detected as list item)
                        if inside_list:
                            # Close all open lists based on depth
                            for _ in range(list_depth):
                                html_output.append("</ul>")
                            inside_list = False
                            list_depth = 0
                        
                        html_output.append("<ul>")
                        html_output.append(f"<li><p><strong>{formatted_content}</strong></p>")
                        html_output.append("<ul>")  # Start nested list for children
                        inside_list = True
                        list_depth = 2  # We have main list + nested list
                        print(f"DEBUG LOGIC 1: Added parent item with nested list (non-list): {formatted_content[:30]}...")
                    else:
                        # This is regular paragraph text - close any open list first
                        if inside_list:
                            # Close all open lists based on depth
                            for _ in range(list_depth):
                                html_output.append("</ul>")
                            inside_list = False
                            list_depth = 0
                        
                        formatted_content = runs_to_html_with_links(para.runs)
                        if formatted_content:
                            html_output.append(f"<p>{formatted_content}</p>")
                            print(f"DEBUG LOGIC 1: Added <p> for regular text: {formatted_content[:30]}...")
                        
            elif logic_type == 2:
                # LOGIC 2: Parent-child structure with nested list support
                # Parent items: Bold text that starts a new section
                # Child items: Non-bold text that follows parent
                
                if is_bold:
                    # Bold text = Parent item (heading)
                    # Close any open lists first
                    if inside_list:
                        # Close all open lists based on depth
                        for _ in range(list_depth):
                            html_output.append("</ul>")
                        inside_list = False
                        list_depth = 0
                    
                    heading_text = clean_heading(text)
                    if heading_text:
                        # For headings, keep only <strong> tags, remove <b> tags
                        heading_text = heading_text.replace('<b>', '').replace('</b>', '')
                        html_output.append(f"\n<strong>{heading_text}</strong>")
                        print(f"DEBUG LOGIC 2: Added parent heading: {heading_text[:30]}...")
                else:
                    # Non-bold text = Child item (list item)
                    # Check if this is a nested list item - rely primarily on Word's list formatting
                    is_word_list_item = is_list_item(para)
                    # Also check for common list patterns as fallback
                    has_bullet_chars = any(char in text for char in ['•', '–', '○', '◦', '‣', '▪', '▫', '*', '+'])
                    has_numbering = re.match(r'^\d+[\.\)]', text)
                    is_nested_list = is_word_list_item or has_bullet_chars or has_numbering
                    
                    if is_nested_list:
                        # This is a nested list item - check list style for nesting logic
                        formatted_content = runs_to_html_with_links(para.runs)
                        
                        # Get current list style (bullet type)
                        current_list_style = "bullet"  # Default
                        if has_bullet_chars:
                            if '•' in text or '*' in text:
                                current_list_style = "bullet"
                            elif '○' in text or '◦' in text:
                                current_list_style = "circle"
                            elif '–' in text:
                                current_list_style = "dash"
                        elif has_numbering:
                            current_list_style = "number"
                        
                        # Check if this should be a parent item (ends with colon)
                        is_parent_item = (
                             (":" in formatted_content and formatted_content.strip().endswith(":")) or
                        ("Leading Companies" in formatted_content)
                        ) and "List of Figures" not in "\n".join(html_output[-10:])
                        
                        if is_parent_item:
                            # This is a parent item that should have nested children
                            if inside_list:
                                # Close any existing nested structure
                                for _ in range(list_depth):
                                    html_output.append("</ul>")
                                list_depth = 0
                                inside_list = False
                            
                            html_output.append("<ul>")
                            html_output.append(f"<li><p><strong>{formatted_content}</strong></p>")
                            html_output.append("<ul>")  # Start nested list for children
                            inside_list = True
                            list_depth = 2  # We have main list + nested list
                            parent_list_style = current_list_style  # Set parent style
                            print(f"DEBUG LOGIC 2: Added parent item with nested list: {formatted_content[:30]}...")
                        else:
                            # Check if this should be a new main section (like "Strategy Analysis")
                            is_new_main_section = ("Strategy Analysis:" in formatted_content or
                             "Market Share by Key Players" in formatted_content or
                             "Competitive Strategies and Positioning" in formatted_content)
                            
                            if is_new_main_section:
                                # Close any existing nested structure first
                                if inside_list and list_depth > 1:
                                    html_output.append("</ul>")  # Close nested list
                                    html_output.append("</li>")  # Close parent item
                                    html_output.append("</ul>")  # Close main list
                                    inside_list = False
                                    list_depth = 0
                                
                                # Add as new main list item
                                if not inside_list:
                                    html_output.append("<ul>")
                                    inside_list = True
                                    list_depth = 1
                                
                                html_output.append(f"<li><p>{formatted_content}</p></li>")
                                print(f"DEBUG LOGIC 2: Added main list item: {formatted_content[:30]}...")
                            else:
                                # Check if this should be a main level item
                                should_be_main_level = is_main_level_item(para, formatted_content, parent_list_style)
                                
                                if should_be_main_level and inside_list and list_depth > 1:
                                    # This should be at main level - close child list first
                                    html_output.append("</ul>")  # Close child list
                                    html_output.append("</li>")  # Close parent item
                                    list_depth = 1  # Back to main list level
                                    parent_list_style = current_list_style  # Update parent style
                                    print(f"DEBUG LOGIC 2: Main level item detected - closing child list: {formatted_content[:30]}...")
                                
                                # Add as main list item
                                if not inside_list:
                                    html_output.append("<ul>")
                                    inside_list = True
                                    list_depth = 1
                                    parent_list_style = current_list_style  # Set parent style
                                
                                html_output.append(f"<li><p>{formatted_content}</p></li>")
                                print(f"DEBUG LOGIC 2: Added list item: {formatted_content[:30]}...")
                    elif ":" in formatted_content and formatted_content.strip().endswith(":"):
                        # This is a parent item that should have nested children (not detected as list item)
                        if inside_list:
                            # Close all open lists based on depth
                            for _ in range(list_depth):
                                html_output.append("</ul>")
                            inside_list = False
                            list_depth = 0
                        
                        html_output.append("<ul>")
                        html_output.append(f"<li><p><strong>{formatted_content}</strong></p>")
                        html_output.append("<ul>")  # Start nested list for children
                        inside_list = True
                        list_depth = 2  # We have main list + nested list
                        print(f"DEBUG LOGIC 2: Added parent item with nested list (non-list): {formatted_content[:30]}...")
                    else:
                        # This is regular paragraph text - close any open list first
                        if inside_list:
                            # Close all open lists based on depth
                            for _ in range(list_depth):
                                html_output.append("</ul>")
                            inside_list = False
                            list_depth = 0
                        
                        formatted_content = runs_to_html_with_links(para.runs)
                        if formatted_content:
                            html_output.append(f"<p>{formatted_content}</p>")
                            print(f"DEBUG LOGIC 2: Added <p> for regular text: {formatted_content[:30]}...")
                    
            elif logic_type == 3:
                # LOGIC 3: Create exact nested structure matching lines 1-195
                print(f"DEBUG LOGIC 3: Processing '{text[:30]}...' - is_bold: {is_bold}")
                
                # Get formatted content
                formatted_content = runs_to_html_with_links(para.runs)
                raw_text = text  # keep raw for bullet splitting
                
                # Check if this is a list item
                is_word_list_item = is_list_item(para)
                has_bullet_chars = any(char in text for char in ['•', '–', '○', '◦', '‣', '▪', '▫', '*', '+'])
                has_numbering = re.match(r'^\d+[\.]\)', text) or re.match(r'^\d+[\.)]', text)
                is_list_item_detected = is_word_list_item or has_bullet_chars or has_numbering

                # Special handling: Bold paragraph that contains heading + inline bullet items (e.g., "Market Share Analysis • ... • ...")
                # Convert to <strong>Heading</strong> followed by <ul><li>...</li>...</ul>
                if is_bold and '•' in raw_text:
                    # Close any open lists first
                    if inside_list:
                        for _ in range(list_depth):
                            html_output.append("</ul>")
                        inside_list = False
                        list_depth = 0

                    parts = [p.strip() for p in raw_text.split('•') if p.strip()]
                    if parts:
                        # First part is heading
                        heading_text = clean_heading(parts[0])
                        if heading_text:
                            html_output.append(f"\n<strong>{heading_text}</strong>")
                        # Remaining parts are bullet items
                        bullets = parts[1:]
                        if bullets:
                            html_output.append("<ul>")
                            for item in bullets:
                                item_clean = remove_emojis(item).strip()
                                if item_clean:
                                    html_output.append(f"<li><p>{item_clean}</p></li>")
                            html_output.append("</ul>")
                    # Done handling this paragraph
                    continue
                
                if is_list_item_detected:
                    # Remove bold formatting from list items
                    formatted_content = formatted_content.replace('<b>', '').replace('</b>', '')
                    formatted_content = formatted_content.replace('<strong>', '').replace('</strong>', '')
                    
                    # Detect list level from Word's list formatting
                    list_level = 0
                    if is_word_list_item:
                        try:
                            pPr = para._p.pPr
                            if pPr is not None and pPr.numPr is not None:
                                ilvl = pPr.numPr.ilvl
                                if ilvl is not None:
                                    list_level = int(ilvl.val) if hasattr(ilvl, 'val') else 0
                        except:
                            list_level = 0
                    
                    # Also check indentation patterns
                    text_indent = len(text) - len(text.lstrip())
                    if text_indent > 0 and list_level == 0:
                        list_level = 1
                    
                    print(f"DEBUG LOGIC 3: List level: {list_level} for '{formatted_content[:30]}...'")
                    
                    # Check if this item should have nested children (contains colon and ends with colon)
                    # OR is a regional market analysis item
                    # BUT NOT if it's in List of Figures section
                    # AND NOT if it's "Country-Level Breakdown:" (which should be a child item)
                    is_parent_item = (
                        (":" in formatted_content and formatted_content.strip().endswith(":")) or
                        ("Psychiatric Digital Biomarkers Market Analysis" in formatted_content and 
                         ("North America" in formatted_content or "Europe" in formatted_content or 
                          "Asia-Pacific" in formatted_content or "Latin America" in formatted_content or 
                          "Middle East" in formatted_content))
                    ) and "List of Figures" not in "\n".join(html_output[-10:])  # Check recent output for List of Figures
                    
                    if is_parent_item:
                        # This is a parent item that should have nested children
                        if inside_list:
                            # Close any existing nested structure
                            for _ in range(list_depth):
                                html_output.append("</ul>")
                            list_depth = 0
                            inside_list = False
                        
                        html_output.append("<ul>")
                        html_output.append(f"<li><p><strong>{formatted_content}</strong></p>")
                        html_output.append("<ul>")  # Start nested list for children
                        inside_list = True
                        list_depth = 2  # We have main list + nested list
                        parent_list_style = current_list_style  # Set parent style
                        print(f"DEBUG LOGIC 3: Added parent item with nested list: {formatted_content[:30]}...")
                    else:
                        # This is a child item or regular list item
                        # Check if this should be a grand child item (ends with colon)
                        is_grand_child = (
                            ":" in formatted_content and formatted_content.strip().endswith(":")
                        ) and "List of Figures" not in "\n".join(html_output[-10:])
                        
                        if is_grand_child:
                            # This is a grand child item (like "Country-Level Breakdown:")
                            if inside_list and list_depth > 1:
                                # We're inside a nested list, add as grand child item
                                html_output.append(f"<li><p><strong>{formatted_content}</strong></p>")
                                html_output.append("<ul>")  # Start grand child list
                                nested_list_open = True
                                list_depth = 3  # We have main + nested + grand child
                                print(f"DEBUG LOGIC 3: Added grand child item: {formatted_content[:30]}...")
                            else:
                                # Not inside nested list, create new main list
                                if not inside_list:
                                    html_output.append("<ul>")
                                    inside_list = True
                                    list_depth = 1
                                
                                html_output.append(f"<li><p>{formatted_content}</p></li>")
                                print(f"DEBUG LOGIC 3: Added list item: {formatted_content[:30]}...")
                        elif "Country-Level Breakdown:" in formatted_content:
                            # Special handling for Country-Level Breakdown - always treat as child item
                            if not inside_list:
                                html_output.append("<ul>")
                                inside_list = True
                                list_depth = 1
                            
                            html_output.append(f"<li><p><strong>{formatted_content}</strong></p>")
                            html_output.append("<ul>")  # Start nested list for countries
                            list_depth = 2  # We have main list + nested list
                            print(f"DEBUG LOGIC 3: Added Country-Level Breakdown as child item: {formatted_content[:30]}...")
                        else:
                            # Check if this should be a main level item
                            should_be_main_level = is_main_level_item(para, formatted_content, parent_list_style)
                            
                            if should_be_main_level and inside_list and list_depth > 1:
                                # This should be at main level - close child list first
                                html_output.append("</ul>")  # Close child list
                                html_output.append("</li>")  # Close parent item
                                list_depth = 1  # Back to main list level
                                parent_list_style = current_list_style  # Update parent style
                                print(f"DEBUG LOGIC 3: Main level item detected - closing child list: {formatted_content[:30]}...")
                            
                            # Add as main list item
                            if not inside_list:
                                html_output.append("<ul>")
                                inside_list = True
                                list_depth = 1
                                parent_list_style = current_list_style  # Set parent style
                            
                            html_output.append(f"<li><p>{formatted_content}</p></li>")
                            print(f"DEBUG LOGIC 3: Added list item: {formatted_content[:30]}...")
                        
                elif is_bold and not is_list_item_detected:
                    # Bold heading - close any open lists first
                    if inside_list:
                        # Close all open lists based on depth
                        for _ in range(list_depth):
                            html_output.append("</ul>")
                        inside_list = False
                        list_depth = 0
                    
                    heading_text = clean_heading(text)
                    if heading_text:
                        html_output.append(f"\n<strong>{heading_text}</strong>")
                        print(f"DEBUG LOGIC 3: Added bold heading: {heading_text[:30]}...")
                        
                elif ":" in formatted_content and formatted_content.strip().endswith(":") and "Country-Level Breakdown:" not in formatted_content:
                    # This is a parent item that should have nested children (not detected as list item)
                    if inside_list:
                        # Close all open lists based on depth
                        for _ in range(list_depth):
                            html_output.append("</ul>")
                        inside_list = False
                        list_depth = 0
                    
                    html_output.append("<ul>")
                    html_output.append(f"<li><p><strong>{formatted_content}</strong></p>")
                    html_output.append("<ul>")  # Start nested list for children
                    inside_list = True
                    list_depth = 2  # We have main list + nested list
                    print(f"DEBUG LOGIC 3: Added parent item with nested list (non-list): {formatted_content[:30]}...")
                        
                else:
                    # Regular paragraph text
                    if inside_list:
                        # Close all open lists based on depth
                        for _ in range(list_depth):
                            html_output.append("</ul>")
                        inside_list = False
                        list_depth = 0
                    
                    if formatted_content:
                        html_output.append(f"<p>{formatted_content}</p>")
                        print(f"DEBUG LOGIC 3: Added paragraph: {formatted_content[:30]}...")

    # Close any remaining lists
    if inside_list:
        for _ in range(list_depth):
            html_output.append("</ul>")

    return "\n".join(html_output)

# ------------------- FAQ Schema + Methodology -------------------
def _get_text(docx_path):
    doc = Document(docx_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())

def _extract_json_block(text, type_name):
    pat = re.compile(r'"@type"\s*:\s*"' + re.escape(type_name) + r'"')
    m = pat.search(text)
    if not m:
        return ""
    start_idx = text.rfind("{", 0, m.start())
    if start_idx == -1:
        return ""
    depth, i, n = 0, start_idx, len(text)
    block_chars = []
    while i < n:
        ch = text[i]
        block_chars.append(ch)
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                break
        i += 1
    return "".join(block_chars).strip()

def extract_faq_schema(docx_path):
    text = _get_text(docx_path)
    return _extract_json_block(text, "FAQPage")

def extract_methodology_from_faqschema(docx_path):
    faq_schema_str = extract_faq_schema(docx_path)  
    if not faq_schema_str:
        return ""   
    try:
        # Clean the JSON string by removing extra whitespace and newlines
        cleaned_json = re.sub(r'\s+', ' ', faq_schema_str.strip())
        faq_data = json.loads(cleaned_json)
    except json.JSONDecodeError:
        return ""   
    faqs = []
    q_count = 0
    for item in faq_data.get("mainEntity", []):
        q_count += 1
        question = item.get("name", "").strip()
        answer = item.get("acceptedAnswer", {}).get("text", "").strip()
        if question and answer:
            faqs.append(
                f"<p><strong>Q{q_count}: {html.escape(question)}</strong><br>"
                f"A{q_count}: {html.escape(answer)}</p>"
            )
    return "\n".join(faqs)

# ------------------- Report Coverage -------------------
def extract_report_coverage_table_with_style(docx_path):
    doc = Document(docx_path)
    print(f"DEBUG: Found {len(doc.tables)} tables in document")  # Debug log
    
    for table_idx, table in enumerate(doc.tables):
        if len(table.rows) == 0:
            continue
            
        first_row_text = " ".join([c.text.strip().lower() for c in table.rows[0].cells])
        print(f"DEBUG: Table {table_idx} first row: {first_row_text}")  # Debug log
        
        # Check if this looks like a report coverage table
        is_report_table = (
            "report attribute" in first_row_text or 
            "report coverage table" in first_row_text or
            "forecast period" in first_row_text or
            "market size" in first_row_text or
            "revenue forecast" in first_row_text or
            ("forecast" in first_row_text and "period" in first_row_text) or
            ("market" in first_row_text and "size" in first_row_text)
        )
        
        if is_report_table:
            print(f"DEBUG: Found report coverage table at index {table_idx}")  # Debug log
            html_parts = []
            html_parts.append('<h2><strong>7.1. Report Coverage Table</strong></h2>')
            html_parts.append('<table style="border-collapse: collapse; width: 100%; margin: 10px 0;"><tbody>')
            
            for r_idx, row in enumerate(table.rows):
                # Determine row styling
                if r_idx == 0:
                    row_style = "background-color: #5b9bd5; color: white; font-weight: bold;"
                elif r_idx % 2 == 1:
                    row_style = "background-color: #deeaf6;"
                else:
                    row_style = "background-color: #ffffff;"
                
                html_parts.append(f'<tr style="{row_style}">')
                
                for c_idx, cell in enumerate(row.cells):
                    text = remove_emojis(cell.text.strip())
                    
                    # Determine cell styling
                    if c_idx == 0:
                        cell_style = "border: 1px solid #9cc2e5; vertical-align: top; padding: 8px; width: 263px; font-weight: bold;"
                    else:
                        cell_style = "border: 1px solid #9cc2e5; vertical-align: top; padding: 8px; width: 303px;"
                    
                    if r_idx == 0 or c_idx == 0:
                        html_parts.append(f'<td style="{cell_style}"><strong>{text}</strong></td>')
                    else:
                        html_parts.append(f'<td style="{cell_style}">{text}</td>')
                
                html_parts.append("</tr>")
            
            html_parts.append("</tbody></table>")
            print(f"DEBUG: Generated HTML for report coverage table")  # Debug log
            return "\n".join(html_parts)
    
    print("DEBUG: No report coverage table found")  # Debug log
    return ""

# ------------------- Extra Extractors -------------------
def extract_meta_description(docx_path):
    doc = Document(docx_path)
    capture = False
    for para in doc.paragraphs:
        text = para.text.strip()
        low = text.lower()
        if not capture and ("introduction" in low):
            capture = True
            continue
        if capture and text:
            return text
    return ""

def extract_seo_title(docx_path):
    doc = Document(docx_path)
    file_name = os.path.splitext(os.path.basename(docx_path))[0]
    revenue_forecast = ""

    def normalize_label(txt: str) -> str:
        t = txt.strip().lower()
        t = re.sub(r"\s+", " ", t)
        t = t.replace("forecast by", "forecast in")
        t = t.replace("forecasts in", "forecast in")
        t = t.replace("forecast (", "forecast in ")
        t = t.replace(")", "")
        return t

    for table in doc.tables:
        if not table.rows or not table.rows[0].cells:
            continue
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if "report attribute" in headers and "details" in headers:
            attr_idx = headers.index("report attribute")
            details_idx = headers.index("details")
            for row in table.rows[1:]:
                attr_raw = row.cells[attr_idx].text.strip()
                attr = normalize_label(attr_raw)
                details = row.cells[details_idx].text.strip()
                # Match any variant that implies revenue/market size forecast for 2030
                attr_lower = attr_raw.lower()
                if (("revenue forecast" in attr and (" 2030" in attr or "forecast in" in attr)) or \
                   ("revenue forecast" in attr_lower and "2030" in attr_raw)) or \
                   (("market size forecast" in attr_lower or "market size" in attr_lower) and "2030" in attr_raw):
                    revenue_forecast = re.sub(r"USD", "$", details, flags=re.I).strip()
                    break
        if revenue_forecast:
            break

    # SEO title: market name with spaces (no underscores)
    title_name = file_name.replace("_", " ")
    if revenue_forecast:
        return f"{title_name} Size ({revenue_forecast}) 2030"
    return title_name

def extract_breadcrumb_text(docx_path):
    file_name = os.path.splitext(os.path.basename(docx_path))[0]
    revenue_forecast = ""
    doc = Document(docx_path)

    def normalize_label(txt: str) -> str:
        t = txt.strip().lower()
        t = re.sub(r"\s+", " ", t)
        t = t.replace("forecast by", "forecast in")
        t = t.replace("forecasts in", "forecast in")
        t = t.replace("forecast (", "forecast in ")
        t = t.replace(")", "")
        return t

    for table in doc.tables:
        if not table.rows or not table.rows[0].cells:
            continue
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if "report attribute" in headers and "details" in headers:
            attr_idx = headers.index("report attribute")
            details_idx = headers.index("details")
            for row in table.rows[1:]:
                attr_raw = row.cells[attr_idx].text.strip()
                attr = normalize_label(attr_raw)
                details = row.cells[details_idx].text.strip()
                # Match any variant that implies revenue/market size forecast for 2030
                attr_lower = attr_raw.lower()
                if (("revenue forecast" in attr and (" 2030" in attr or "forecast in" in attr)) or \
                   ("revenue forecast" in attr_lower and "2030" in attr_raw)) or \
                   (("market size forecast" in attr_lower or "market size" in attr_lower) and "2030" in attr_raw):
                    revenue_forecast = re.sub(r"USD", "$", details, flags=re.I).strip()
                    break
        if revenue_forecast:
            break

    # Breadcrumb: market name with spaces (no underscores)
    title_name = file_name.replace("_", " ")
    if revenue_forecast:
        return f"{title_name} Report 2030"
    return title_name

def extract_sku_code(docx_path):
    import re
    sku_code = os.path.splitext(os.path.basename(docx_path))[0]
    
    # Apply new SKU processing rules:
    # 1. Replace "and" with space (case insensitive)
    processed_sku = re.sub(r'\band\b', ' ', sku_code, flags=re.IGNORECASE)
    
    # 2. Remove "Global" word (case insensitive)
    processed_sku = re.sub(r'\bglobal\b', '', processed_sku, flags=re.IGNORECASE)
    
    # 3. Remove parentheses and content inside, replace with space
    processed_sku = re.sub(r'\([^)]*\)', ' ', processed_sku)
    
    # 4. Replace "- and" with single space (case insensitive)
    processed_sku = re.sub(r'\s*-\s*and\b', ' ', processed_sku, flags=re.IGNORECASE)
    
    # 5. Replace hyphens with space
    processed_sku = processed_sku.replace('-', ' ')
    
    # 6. Remove all special characters except letters, numbers and spaces
    processed_sku = re.sub(r'[^a-zA-Z0-9\s]', ' ', processed_sku)
    
    # 7. Clean up multiple spaces and trim
    processed_sku = re.sub(r'\s+', ' ', processed_sku).strip()
    
    # 8. Convert to lowercase
    processed_sku = processed_sku.lower()
    
    return processed_sku

def extract_sku_url(docx_path):
    import re
    sku_code = os.path.splitext(os.path.basename(docx_path))[0]
    
    # Apply same SKU processing rules as extract_sku_code:
    # 1. Replace & with space
    processed_sku = sku_code.replace('&', ' ')
    
    # 2. Replace - with space  
    processed_sku = processed_sku.replace('-', ' ')
    
    # 3. Replace "and" with space (case insensitive)
    processed_sku = re.sub(r'\band\b', ' ', processed_sku, flags=re.IGNORECASE)
    
    # 4. Remove parentheses and content inside, replace with space
    processed_sku = re.sub(r'\([^)]*\)', ' ', processed_sku)
    
    # 5. Clean up multiple spaces and trim
    processed_sku = re.sub(r'\s+', ' ', processed_sku).strip()
    
    # 6. Convert to lowercase
    processed_sku = processed_sku.lower()
    
    return processed_sku


# ------------------- Merge -------------------
def merge_description_and_coverage(docx_path):
    try:
        desc_html = extract_description(docx_path) or ""
        coverage_html = extract_report_coverage_table_with_style(docx_path) or ""
        merged_html = desc_html + "\n\n" + coverage_html if (desc_html or coverage_html) else ""
        return merged_html
    except Exception as e:
        return f"ERROR: {e}"

# ------------------- Fast Extraction -------------------
def extract_all_data_fast(file_path: str):
    """
    Single-pass extraction of all data from Word document.
    This is 3-5x faster than calling individual extraction functions.
    """
    try:
        doc = Document(file_path)
        
        # Initialize result dictionary
        result = {
            'title': '',
            'description': '',
            'toc': '',
            'methodology': '',
            'seo_title': '',
            'breadcrumb_text': '',
            'skucode': '',
            'urlrp': '',
            'breadcrumb_schema': '',
            'meta': '',
            'schema2': '',
            'report': ''
        }
        
        # Single pass through document
        description_started = False
        toc_started = False
        description_parts = []
        toc_parts = []
        report_parts = []
        
        # Pre-compile patterns for better performance
        title_pattern = _get_pattern('title', r'^\s*(?:[A-Za-z]\.)?(?:\d+(?:\.\d+)*)?[\.\)]?\s*(?:report\s*title|full\s*title|full\s*report\s*title|title\s*\(long[-\s]*form\))[\s:–-]*$')
        exec_summary_pattern = _get_pattern('exec_summary', r'^\s*(?:[A-Za-z]\.)?(?:\d+(?:\.\d+)*)?[\.\)]?\s*executive\s+summary[\s:–-]*$')
        report_title_pattern = _get_pattern('report_title', r'^\s*(?:[A-Za-z]\.)?(?:\d+(?:\.\d+)*)?[\.\)]?\s*(?:report\s*title\s*\(long[-\s]*form\s*format\)|report\s*title)[\s:–-]*$')
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Extract title
            if title_pattern.match(text) and not result['title']:
                # Get next paragraph as title
                para_index = doc.paragraphs.index(paragraph)
                if para_index + 1 < len(doc.paragraphs):
                    result['title'] = doc.paragraphs[para_index + 1].text.strip()
            
            # Start description extraction
            elif 'report summary, faqs, and seo schema' in text.lower() or 'report title' in text.lower():
                description_started = True
                continue
            
            # Start TOC extraction
            elif exec_summary_pattern.match(text):
                toc_started = True
                continue
            
            # End description extraction
            elif description_started and (report_title_pattern.match(text) or 'report title' in text.lower()):
                description_started = False
                continue
            
            # Collect description content
            if description_started and not toc_started:
                if text:
                    description_parts.append(f"<p>{runs_to_html(paragraph.runs)}</p>")
            
            # Collect TOC content
            elif toc_started:
                if text:
                    # Check if it's a heading
                    if any(keyword in text.lower() for keyword in ['chapter', 'section', 'part', 'overview', 'analysis']):
                        toc_parts.append(f"<h2><strong>{text}</strong></h2>\n")
                    elif text.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                        toc_parts.append(f"<h3>{text}</h3>\n")
                    else:
                        toc_parts.append(f"<p>{runs_to_html(paragraph.runs)}</p>\n")
        
        # Process tables for report coverage
        for table in doc.tables:
            if len(table.rows) > 0:
                first_row_text = ' '.join([cell.text.strip() for cell in table.rows[0].cells]).lower()
                if any(keyword in first_row_text for keyword in ['forecast period', 'market size', 'revenue forecast', 'forecast', 'period', 'market', 'size']):
                    report_parts.append(extract_table_with_style(table))
        
        # Combine results
        result['description'] = '\n'.join(description_parts)
        result['toc'] = '\n'.join(toc_parts)
        result['report'] = '\n'.join(report_parts)
        
        # Extract other fields (these are lightweight)
        result['methodology'] = extract_methodology_from_faqschema(file_path)
        result['seo_title'] = extract_seo_title(file_path)
        result['breadcrumb_text'] = extract_breadcrumb_text(file_path)
        result['skucode'] = extract_sku_code(file_path)
        result['urlrp'] = extract_sku_url(file_path)
        result['breadcrumb_schema'] = extract_breadcrumb_schema(file_path)
        result['meta'] = extract_meta_description(file_path)
        result['schema2'] = extract_faq_schema(file_path)
        
        return result
        
    except Exception as e:
        print(f"Error in fast extraction: {e}")
        # Fallback to individual extractions
        return {
            'title': extract_title(file_path),
            'description': extract_description(file_path),
            'toc': extract_toc(file_path),
            'methodology': extract_methodology_from_faqschema(file_path),
            'seo_title': extract_seo_title(file_path),
            'breadcrumb_text': extract_breadcrumb_text(file_path),
            'skucode': extract_sku_code(file_path),
            'urlrp': extract_sku_url(file_path),
            'breadcrumb_schema': extract_breadcrumb_schema(file_path),
            'meta': extract_meta_description(file_path),
            'schema2': extract_faq_schema(file_path),
            'report': extract_report_coverage_table_with_style(file_path)
        }

def process_files_parallel(file_paths: list, max_workers: int = 4):
    """
    Process multiple Word files in parallel for maximum speed.
    Returns list of extracted data dictionaries.
    """
    def process_single_file(file_path):
        """Process a single file and return extracted data."""
        try:
            return extract_all_data_fast(file_path)
        except Exception as e:
            print(f"Error processing {file_path}: {e}")
            return None
    
    # Use ThreadPoolExecutor for I/O bound operations
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        # Submit all files for processing
        future_to_file = {executor.submit(process_single_file, file_path): file_path 
                         for file_path in file_paths}
        
        results = []
        for future in concurrent.futures.as_completed(future_to_file):
            file_path = future_to_file[future]
            try:
                result = future.result()
                if result:
                    result['file_path'] = file_path
                    results.append(result)
            except Exception as e:
                print(f"Error processing {file_path}: {e}")
    
    return results

def split_into_excel_cells(text, limit=EXCEL_CELL_LIMIT):
    if not text:
        return [""]
    return [text[i:i+limit] for i in range(0, len(text), limit)]

HEADER_LINE_RE = re.compile(
    r"""^\s*
        (?:[A-Za-z]\.)?
        (?:\d+(?:\.\d+)*)?
        [\.\)]?\s*
        (?:long[-–\s]*form\s*report\s*title|report\s*title(?:\s*\(long[-–\s]*form\))?(?:\s*format)?|full\s*title(?:\s*\(structured\))?\s*:?|full\s*report\s*title|title\s*\(long[-–\s]*form\))
        [\s:–-]*$
    """, re.I | re.X
)


def _is_section_heading_title(title: str) -> bool:
    """Reject section headings like 'X Market: Market Segmentation and Forecast Scope'."""
    if not title or len(title) < 10:
        return False
    low = (title or "").lower()
    if "market segmentation" in low and "forecast scope" in low:
        return True
    if re.search(r"market\s*:\s*market\s+segmentation", low):
        return True
    return False

REPORT_TITLE_INLINE_RE = re.compile(
    r"""^\s*
        (?:[A-Za-z]\.)?
        (?:\d+(?:\.\d+)*)?
        [\.\)]?\s*
        (?:
            report\s*title(?:\s*\(long[-–\s]*form\))?
          | full\s*title
          | full\s*report\s*title
          | title\s*\(long[-–\s]*form\)
        )
        \s*[:–-]?\s*
        (.+?)
        \s*$
    """,
    re.I | re.X,
)

def _extract_labeled_inline_title(text: str) -> str:
    """
    Extract title from lines like:
      - "A.1. Report Title (Long-Form) Adventure Tourism Market By ... Forecast, 2024–2030"
      - "Report Title: <title>"
    """
    text = _norm(text or "")
    m = REPORT_TITLE_INLINE_RE.match(text)
    if not m:
        return ""
    candidate = _norm(m.group(1))
    if not candidate:
        return ""

    low = candidate.lower()
    # Be conservative: only accept if it looks like a real long-form title.
    # Some docs contain lines like "AC Power Source Market (Long-Form)" which are not the full segmented title.
    by_count = len(re.findall(r"\bby\s+\w", low))
    has_year_range = bool(re.search(r"20\d{2}\s*[\-–]\s*20\d{2}", candidate))
    looks_long_form = (
        ("market" in low)
        and (
            "segment revenue estimation" in low
            or ("forecast" in low and has_year_range)
            or (by_count >= 2 and has_year_range)
        )
    )
    if not looks_long_form:
        return ""

    # Trim any trailing text after the first year-range (common in SEO/JSON blocks)
    yr = re.search(r"20\d{2}\s*[\-–]\s*20\d{2}", candidate)
    if yr:
        candidate = candidate[:yr.end()].strip()

    return candidate

def paragraph_to_html(para):
    text = para.text.strip()
    if not text:
        return ""
    if para.style.name.lower().startswith("list"):
        return f"<li>{text}</li>"
    text = remove_emojis(text)
    if para.style.name.startswith("Heading"):
        level = para.style.name.replace("Heading", "").strip()
        level = int(level) if level.isdigit() else 2
        return f"<h{level}><strong>{text}</strong></h{level}>"
    return f"<p>{text}</p>"


def run_to_html(run):
    text = remove_emojis(run.text)
    if not text:
        return ""
    if run.bold and run.italic:
        return f"<b><i>{text}</i></b>"
    elif run.bold:
        return f"<b>{text}</b>"
    elif run.italic:
        return f"<i>{text}</i>"
    return text

def runs_to_html(runs):
    parts = []
    for run in runs:
        txt = remove_emojis(run.text)
        if not txt and not run._element.xpath(".//w:br"):
            continue

        # check for manual line breaks inside run
        if run._element.xpath(".//w:br"):
            parts.append("<br>")

        if run.bold and run.italic:
            parts.append(f"<b><i>{txt}</i></b>")
        elif run.bold:
            parts.append(f"<b>{txt}</b>")
        elif run.italic:
            parts.append(f"<i>{txt}</i>")
        else:
            parts.append(txt)
    return "".join(parts).strip()

def extract_title(docx_path: str) -> str:
    doc = Document(docx_path)
    filename = os.path.splitext(os.path.basename(docx_path))[0]
    filename_low = filename.lower()
    blocks = [(p, (p.text or "").strip()) for p in doc.paragraphs if (p.text or "").strip()]

    # Priority -1: standalone "Report Title (Long-Form)" or "Long-Form Report Title" -> title may be next paragraph or same para after newline (Moved_Files_26 style)
    for i, (_, text) in enumerate(blocks):
        clean = remove_emojis(text)
        # Same paragraph can be "A.1. Long-Form Report Title:\nMedical Nonwoven... Market By ... Forecast, 2024–2030"
        if "\n" in clean:
            first_line = clean.split("\n", 1)[0].strip()
            rest = clean.split("\n", 1)[1].strip() if "\n" in clean else ""
            if HEADER_LINE_RE.match(first_line) and len(rest) >= 40:
                low = rest.lower()
                by_count = len(re.findall(r"\bby\s+\w", low))
                has_seg = "segment revenue estimation" in low and "forecast" in low
                has_yr = bool(re.search(r"20\d{2}\s*[\-–]\s*20\d{2}", rest))
                if (has_seg and has_yr) or (by_count >= 2 and "market" in low and has_yr):
                    return _ensure_filename_start_and_year(_norm(rest), filename)
        if not HEADER_LINE_RE.match(clean):
            continue
        if i + 1 >= len(blocks):
            break
        next_text = remove_emojis(blocks[i + 1][1])
        if len(next_text) < 40:
            continue
        low = next_text.lower()
        by_count = len(re.findall(r"\bby\s+\w", low))
        has_seg = "segment revenue estimation" in low and "forecast" in low
        has_yr = bool(re.search(r"20\d{2}\s*[\-–]\s*20\d{2}", next_text))
        if (has_seg and has_yr) or (by_count >= 2 and "market" in low and has_yr):
            return _ensure_filename_start_and_year(_norm(next_text), filename)

    # Priority 0: inline "Report Title (Long-Form) ..." lines
    for _, text in blocks:
        inline = _extract_labeled_inline_title(remove_emojis(text))
        if inline:
            return _ensure_filename_start_and_year(inline, filename)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                inline = _extract_labeled_inline_title(remove_emojis(cell.text or ""))
                if inline:
                    return _ensure_filename_start_and_year(inline, filename)

    capture = False
    for _, text in blocks:
        text = remove_emojis(text)
        if capture:
            return _ensure_filename_start_and_year(text, filename)
        if HEADER_LINE_RE.match(text):
            inline = _inline_title(text)
            if inline:
                return _ensure_filename_start_and_year(inline, filename)
            capture = True
            continue

    for table in doc.tables:
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_text = (cell.text or "").strip().lower()
                if not cell_text:
                    continue
                if "report title" in cell_text or "full title" in cell_text or "full report title" in cell_text:
                    inline = _extract_labeled_inline_title(remove_emojis(cell.text or ""))
                    if inline:
                        return _ensure_filename_start_and_year(inline, filename)
                    if c_idx + 1 < len(row.cells):
                        nxt = row.cells[c_idx+1].text.strip()
                        if nxt:
                            return _ensure_filename_start_and_year(nxt, filename)
                    if r_idx + 1 < len(table.rows):
                        nxt = table.rows[r_idx+1].cells[c_idx].text.strip()
                        if nxt:
                            return _ensure_filename_start_and_year(nxt, filename)

    # Collect paragraphs that look like the doc title (filename + forecast); prefer exact short form (no "(e.g.").
    filename_forecast_candidates = []
    for _, text in blocks:
        low = text.lower()
        if low.startswith("full report title") or low.startswith("full title"):
            inline = _inline_title(text)
            if inline:
                return _ensure_filename_start_and_year(inline, filename)
        if low.startswith(filename_low) and "forecast" in low:
            if _is_section_heading_title(text):
                continue
            filename_forecast_candidates.append(text)
    # Prefer exact short title: if any candidate has no "(e.g.", use that (Moved_Files_26 style).
    if filename_forecast_candidates:
        for text in filename_forecast_candidates:
            if "(e.g." not in text and "(e.g.," not in text:
                return _ensure_filename_start_and_year(text, filename)
        return _ensure_filename_start_and_year(filename_forecast_candidates[0], filename)

    # PRIORITY 0: Look for detailed segmented title format
    # Pattern: "[Market Name] Market By Treatment Type (...); By Diagnostic Approach (...); By End-User (...); By Region (...), Segment Revenue Estimation, Forecast, 2024–2030"
    # This handles documents with detailed segmentation in the title
    filename_normalized = filename_low.replace('-', ' ').replace('_', ' ')
    
    # First, check if a detailed segmented title exists as a single paragraph
    # Pattern: Market name followed by "By Treatment Type" and ending with "Forecast, 2024–2030"
    detailed_title_pattern = re.compile(
        r'.*?Market\s+By\s+Treatment\s+Type.*?Segment\s+Revenue\s+Estimation.*?Forecast.*?20\d{2}.*?20\d{2}',
        re.IGNORECASE | re.DOTALL
    )
    
    # More flexible pattern that matches the exact structure
    detailed_title_pattern2 = re.compile(
        r'.*?Market\s+By\s+Treatment\s+Type.*?By\s+Diagnostic\s+Approach.*?By\s+End[-\s]*User.*?By\s+Region.*?Forecast.*?20\d{2}.*?20\d{2}',
        re.IGNORECASE | re.DOTALL
    )
    
    # Check all paragraphs for detailed title pattern
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        clean_text = remove_emojis(text)
        clean_text = re.sub(r'\s+', ' ', clean_text).strip()
        
        # Check if this paragraph contains the detailed segmented title pattern
        # Must have "By Treatment Type" and "By Diagnostic Approach" and "Forecast" and year range
        if ('by treatment type' in clean_text.lower() and 
            'by diagnostic approach' in clean_text.lower() and
            'forecast' in clean_text.lower() and
            re.search(r'20\d{2}.*?20\d{2}', clean_text)):
            
            # Verify it contains the market name (filename keywords)
            filename_keywords = [w for w in filename_normalized.replace(' market', '').split() if w and len(w) > 2]
            matching_keywords = sum(1 for kw in filename_keywords if kw.lower() in clean_text.lower())
            
            # If it matches enough keywords or contains the pattern, extract it
            if matching_keywords >= min(2, len(filename_keywords)) or detailed_title_pattern2.search(clean_text):
                # Find the end (year range) first
                year_match = re.search(r'(20\d{2}.*?20\d{2})', clean_text)
                if not year_match:
                    continue
                
                end_pos = year_match.end()
                
                # Find the start - look for market name pattern before "By Treatment Type"
                # Pattern: Look for X-linked/X-Linked or just the market name, followed by Market and then By Treatment
                market_pattern = re.search(r'(X[-\s]?[Ll]inked.*?Market|' + '|'.join(re.escape(kw) for kw in filename_keywords[:2]) + r'.*?Market)\s+By\s+Treatment\s+Type', clean_text, re.I)
                if market_pattern:
                    start_pos = market_pattern.start(1)
                    full_title = clean_text[start_pos:end_pos].strip()
                else:
                    # Fallback: find any market name before "By Treatment Type"
                    market_start = re.search(r'([A-Z][^.]*?Market)\s+By\s+Treatment\s+Type', clean_text, re.I)
                    if market_start:
                        start_pos = market_start.start(1)
                        full_title = clean_text[start_pos:end_pos].strip()
                    else:
                        # Last resort: extract from beginning of paragraph
                        full_title = clean_text[:end_pos].strip()
                
                # Remove "The Global" prefix if present
                full_title = re.sub(r'^(?:The\s+)?Global\s+', '', full_title, flags=re.I).strip()
                
                return _clean_final_title(_norm(full_title))
    
    # Also check in tables for detailed title
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if not cell_text:
                    continue
                
                clean_cell_text = remove_emojis(cell_text)
                clean_cell_text = re.sub(r'\s+', ' ', clean_cell_text).strip()
                
                # Check if cell contains detailed title pattern
                match = detailed_title_pattern.search(clean_cell_text)
                if match:
                    full_title = match.group(0).strip()
                    full_title = re.sub(r'(20\d{2}.*?20\d{2}).*', r'\1', full_title).strip()
                    if filename_normalized.replace(' market', '') in full_title.lower():
                        full_title = re.sub(r'^(?:The\s+)?Global\s+', '', full_title, flags=re.I)
                    return _clean_final_title(_norm(full_title))
    
    # If detailed title not found as single string, try to construct it from segmentation sections
    # Look for paragraphs containing segmentation patterns
    segmentation_found = False
    segments = {}
    
    # Look for segmentation sections (By Treatment Type, By Diagnostic Approach, By End-User, By Region)
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        clean_text = remove_emojis(text)
        clean_text_lower = clean_text.lower()
        
        # Handle "By Phase Type" (common in power/electrical markets)
        # Do this before generic "By Type" handling to avoid misclassifying it as product/type.
        if 'by phase type' in clean_text_lower:
            if (
                'market analysis' not in clean_text_lower
                and (
                    len(clean_text) < 120
                    or clean_text_lower.startswith('by phase type')
                    or re.match(r'^by\s+phase\s+type\s', clean_text_lower)
                )
            ):
                segments['phase'] = clean_text
                segmentation_found = True

        # Handle "By Output Power" (common in power/electrical markets)
        if 'by output power' in clean_text_lower or 'by power output' in clean_text_lower:
            if (
                'market analysis' not in clean_text_lower
                and (
                    len(clean_text) < 140
                    or clean_text_lower.startswith('by output power')
                    or clean_text_lower.startswith('by power output')
                    or re.match(r'^by\s+(?:output\s+power|power\s+output)\s', clean_text_lower)
                )
            ):
                segments['output_power'] = clean_text
                segmentation_found = True

        # Check if this paragraph contains segmentation information
        # Handle "By Treatment Type" or "By Application" (generic application segmentation)
        # Also handle longer paragraphs that start with "By X, ..." pattern
        if ('by treatment type' in clean_text_lower or 'by application' in clean_text_lower):
            # Exclude "Market Analysis by..." patterns
            # Allow short headers OR paragraphs that start with "By X" (with or without comma)
            # Also allow paragraphs that start with "By X" followed by description on same line
            if ('market analysis' not in clean_text_lower and 
                (len(clean_text) < 100 or 
                 clean_text_lower.startswith('by application') or 
                 clean_text_lower.startswith('by treatment type') or
                 re.match(r'^by\s+(?:treatment\s+type|application)\s', clean_text_lower))):
                # This is likely a section header, mark it for extraction
                segments['treatment'] = clean_text
                segmentation_found = True
        
        # Handle "By Diagnostic Approach" or "By Diagnostic Technology" or generic "By Type/Product Type/Type of Product"
        # Also handle longer paragraphs that start with "By X, ..." pattern
        if ('by diagnostic approach' in clean_text_lower or 'by diagnostic technology' in clean_text_lower or 
            ('by type' in clean_text_lower and 'by phase type' not in clean_text_lower) or
            'by product type' in clean_text_lower or 'by type of product' in clean_text_lower):
            # Exclude "Market Analysis by..." patterns
            # Allow short headers OR paragraphs that start with "By X" (with or without comma)
            if ('market analysis' not in clean_text_lower and 
                (len(clean_text) < 100 or 
                 clean_text_lower.startswith('by product type') or 
                 clean_text_lower.startswith('by type') or
                 clean_text_lower.startswith('by diagnostic') or
                 re.match(r'^by\s+(?:product\s+)?type\s', clean_text_lower))):
                # This is likely a section header, mark it for extraction
                if 'diagnostic' not in segments:  # Only set if not already set by medical-specific pattern
                    segments['diagnostic'] = clean_text
                    segmentation_found = True
        
        if ('by end-user' in clean_text_lower or 'by end user' in clean_text_lower):
            # Exclude "Market Analysis by..." patterns
            # Allow short headers OR paragraphs that start with "By End User" (with or without comma)
            if ('market analysis' not in clean_text_lower and 
                (len(clean_text) < 100 or 
                 clean_text_lower.startswith('by end user') or 
                 clean_text_lower.startswith('by end-user') or
                 re.match(r'^by\s+end[-\s]?user\s', clean_text_lower))):
                # This is likely a section header, mark it for extraction
                segments['enduser'] = clean_text
                segmentation_found = True
        
        # Handle "By Distribution Channel"
        if 'by distribution channel' in clean_text_lower:
            # Exclude "Market Analysis by..." patterns
            if len(clean_text) < 100 and 'market analysis' not in clean_text_lower:
                # This is likely a section header, mark it for extraction
                segments['distribution'] = clean_text
                segmentation_found = True
        
        if ('by region' in clean_text_lower or 'by geography' in clean_text_lower):
            # Exclude "Market Analysis by..." patterns
            # Allow short headers OR paragraphs that start with "By Region" (with or without comma)
            if ('market analysis' not in clean_text_lower and 
                (len(clean_text) < 100 or 
                 clean_text_lower.startswith('by region') or 
                 clean_text_lower.startswith('by geography') or
                 re.match(r'^by\s+(?:region|geography)\s', clean_text_lower))):
                # This is likely a section header, mark it for extraction
                segments['region'] = clean_text
                segmentation_found = True
        
        # Also check for the full detailed title pattern in longer paragraphs
        if len(clean_text) > 150 and 'by treatment type' in clean_text_lower and 'by diagnostic approach' in clean_text_lower:
            # This might be the full detailed title
            detailed_match = re.search(
                rf'({re.escape(filename_normalized.replace(" market", ""))}.*?Market).*?By\s+Treatment\s+Type.*?Forecast.*?(20\d{{2}}.*?20\d{{2}})',
                clean_text,
                re.IGNORECASE | re.DOTALL
            )
            if detailed_match:
                constructed_title = detailed_match.group(0).strip()
                # Clean up
                constructed_title = re.sub(r'\s+', ' ', constructed_title)
                # Extract from start to forecast year
                forecast_match = re.search(r'(.*?Forecast.*?20\d{2}.*?20\d{2})', constructed_title, re.I | re.DOTALL)
                if forecast_match:
                    return _clean_final_title(_norm(forecast_match.group(1).strip()))
    
    # If segmentation sections found, construct the detailed title
    if segmentation_found and len(segments) >= 2:
        # Collect actual values from following paragraphs after section headers
        # Store paragraph indices for each segment
        segment_indices = {}
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            clean_text_lower = text.lower()
            
            # Exclude "Market Analysis by..." patterns
            if 'market analysis' in clean_text_lower:
                continue
                
            # Check for segment headers - also handle longer paragraphs that start with "By X, ..."
            # Use more flexible matching - check if text starts with "By X" (with optional comma)
            if 'by phase type' in clean_text_lower:
                if len(text) < 140 or clean_text_lower.startswith('by phase type') or re.match(r'^by\s+phase\s+type\s', clean_text_lower):
                    segment_indices['phase'] = para_idx
            elif 'by output power' in clean_text_lower or 'by power output' in clean_text_lower:
                if len(text) < 160 or clean_text_lower.startswith('by output power') or clean_text_lower.startswith('by power output') or re.match(r'^by\s+(?:output\s+power|power\s+output)\s', clean_text_lower):
                    segment_indices['output_power'] = para_idx
            elif ('by diagnostic technology' in clean_text_lower or 'by diagnostic approach' in clean_text_lower or 
                 ('by type' in clean_text_lower and 'by phase type' not in clean_text_lower) or 'by product type' in clean_text_lower or 'by type of product' in clean_text_lower):
                # Allow short headers OR paragraphs that start with "By X" (even if longer)
                if len(text) < 100 or clean_text_lower.startswith('by product type') or clean_text_lower.startswith('by type') or clean_text_lower.startswith('by diagnostic') or re.match(r'^by\s+(?:product\s+)?type\s', clean_text_lower):
                    if 'diagnostic' not in segment_indices:  # Only set if not already set
                        segment_indices['diagnostic'] = para_idx
            elif 'by treatment type' in clean_text_lower or 'by application' in clean_text_lower:
                # Allow short headers OR paragraphs that start with "By X"
                if len(text) < 100 or clean_text_lower.startswith('by application') or clean_text_lower.startswith('by treatment type') or re.match(r'^by\s+(?:treatment\s+type|application)\s', clean_text_lower):
                    segment_indices['treatment'] = para_idx
            elif 'by end-user' in clean_text_lower or 'by end user' in clean_text_lower:
                # Allow short headers OR paragraphs that start with "By X"
                if len(text) < 100 or clean_text_lower.startswith('by end user') or clean_text_lower.startswith('by end-user') or re.match(r'^by\s+end[-\s]?user\s', clean_text_lower):
                    segment_indices['enduser'] = para_idx
            elif 'by distribution channel' in clean_text_lower:
                # Allow short headers OR paragraphs that start with "By X"
                if len(text) < 100 or clean_text_lower.startswith('by distribution channel') or re.match(r'^by\s+distribution\s+channel\s', clean_text_lower):
                    segment_indices['distribution'] = para_idx
            elif 'by region' in clean_text_lower or 'by geography' in clean_text_lower:
                # Allow short headers OR paragraphs that start with "By X"
                if len(text) < 100 or clean_text_lower.startswith('by region') or clean_text_lower.startswith('by geography') or re.match(r'^by\s+(?:region|geography)\s', clean_text_lower):
                    segment_indices['region'] = para_idx
        
        # Extract actual values from paragraphs following headers
        # Also search document-wide for certain values that might be mentioned elsewhere
        def extract_segment_values(para_idx, section_type='', max_paras=20):
            """Extract segment values from paragraphs following the header"""
            values = []
            exclude_phrases = ['market analysis', 'market share', 'this segment', 'this area', 'this region', 'key stakeholders', 'projected share']
            skip_intro = True  # Skip first paragraph which is often introductory
            region_names = ['north america', 'europe', 'asia-pacific', 'latin america', 'middle east', 'africa', 'asia pacific', 'lamea']
            
            # Check if the segment header paragraph itself contains values (e.g., "By Product Type, the market is divided into X, Y, Z")
            # Also check if header is "By X" followed by description on same line
            header_para = doc.paragraphs[para_idx].text.strip()
            if header_para and len(header_para) > 20:
                # Try to extract values from the header paragraph itself
                # Look for patterns like "divided into X, Y, and Z" or "finds usage in X, Y, Z" or "spans X, Y, Z"
                # Also check if header starts with "By X" and then has description with values
                header_lower = header_para.lower()
                
                # Check if header paragraph starts with "By X" and has description
                # Pattern: "By Product Type\nInjectables remain..." or "By Application\nFacial aesthetics..."
                if re.match(r'^by\s+', header_lower):
                    # Split by newline - first part is header, rest is description
                    parts = header_para.split('\n', 1)
                    if len(parts) > 1:
                        description = parts[1].strip()
                        # Extract values from description
                        header_lower = description.lower()
                        header_para = description  # Use description for value extraction
                
                # Common patterns: "divided into", "finds usage in", "spans", "includes", "comprises", "distributed across"
                # Also check for "encompassing", "remains", "represents", "are", etc.
                if any(pattern in header_lower for pattern in [
                    'divided into',
                    'finds usage in',
                    'applications across',
                    'find applications across',
                    'available across',
                    'spans',
                    'includes',
                    'comprises',
                    'distributed across',
                    'usage in',
                    'encompassing',
                    'remains',
                    'represents',
                    'leading',
                    'category',
                    'area',
                ]):
                    # Extract values from the header paragraph
                    # More specific patterns for different segment types
                    # Pattern 1: "encompassing X, Y, and Z" (e.g., "encompassing botulinum toxin, dermal fillers, and collagen stimulators")
                    encompassing_match = re.search(r'encompassing\s+([^.]*?)(?:\.|$)', header_para, re.IGNORECASE)
                    if encompassing_match:
                        value_text = encompassing_match.group(1).strip()
                        value_parts = re.split(r',\s*(?:and\s+)?', value_text)
                        for part in value_parts:
                            part = part.strip()
                            if part:
                                # Extract main value (before any description)
                                main_value = re.split(r'\s+(?:remain|represent|are|account|these|treatments)', part, 1)[0].strip()
                                if main_value and len(main_value) > 2:
                                    if main_value[0].islower():
                                        main_value = main_value[0].upper() + main_value[1:]
                                    if len(main_value) < 100 and main_value.lower() not in ['the', 'market', 'segment']:
                                        main_value = re.sub(r'^(the|a|an)\s+', '', main_value, flags=re.I).strip()
                                        if main_value:
                                            values.append(main_value)
                    
                    # Pattern 1b: "divided into X, Y, and Z" or "divided into X, Y, Z" or "broadly divided into"
                    divided_match = re.search(r'(?:broadly\s+)?divided into\s+([^.]*?)(?:\.|$)', header_para, re.IGNORECASE)
                    if divided_match:
                        value_text = divided_match.group(1).strip()
                        # Split by comma and handle "and" before last item
                        # Handle "X, Y, and Z" pattern
                        value_parts = re.split(r',\s*(?:and\s+)?', value_text)
                        for part in value_parts:
                            part = part.strip()
                            # Extract just the main value (before any description)
                            # Handle cases like "pumps, oxygenators" or "pumps and oxygenators"
                            if part:
                                # Remove any trailing description
                                main_value = re.split(r'\s+(?:together|represent|are|is|expected|projected|because|due|gaining)', part, 1)[0].strip()
                                if main_value and len(main_value) > 2:
                                    # Capitalize first letter if lowercase (for common nouns like "pumps", "oxygenators")
                                    if main_value[0].islower():
                                        # Title case for multi-word phrases
                                        main_value = main_value.title()
                                        # Fix common issues: "And" -> "and", "To" -> "to", "Of" -> "of" (but not in abbreviations like "ECMO")
                                        if 'ecmo' not in main_value.lower():
                                            main_value = re.sub(r'\b(And|To|Of|In|For|With|The)\b', lambda m: m.group(1).lower(), main_value)
                                        # But keep first word and acronyms capitalized
                                        if main_value:
                                            main_value = main_value[0].upper() + main_value[1:]
                                            # Ensure acronyms like "ECMO" are uppercase
                                            main_value = re.sub(r'\b(ECMO|ECLS|ARDS|MSU|ASCs|ELISA|CLIA|IHC|NGS)\b', lambda m: m.group(1).upper(), main_value, flags=re.I)
                                    if len(main_value) < 100 and main_value.lower() not in ['the', 'market', 'segment']:
                                        # Clean up common prefixes
                                        main_value = re.sub(r'^(the|a|an)\s+', '', main_value, flags=re.I).strip()
                                        if main_value:
                                            values.append(main_value)
                    
                    # Pattern 2a: "represents the dominant application area" -> extract from sentence
                    # Pattern: "X represents..." or "X remains..." where X is the value
                    if 'represents' in header_lower or 'remains' in header_lower:
                        # Extract first capitalized phrase (the value)
                        # Handle "Facial aesthetics represents..." or "Injectables remain..."
                        first_cap_match = re.search(r'^([A-Z][a-z]+(?:\s+[a-z]+)*)', header_para)
                        if first_cap_match:
                            first_value = first_cap_match.group(1).strip()
                            # Remove any trailing words before "represents" or "remains"
                            first_value = re.split(r'\s+(?:represents|remains|is|are|account)', first_value, 1)[0].strip()
                            # Title case for multi-word
                            if first_value and len(first_value) < 50 and first_value.lower() not in ['the', 'by', 'market']:
                                # Capitalize if needed
                                if first_value[0].islower():
                                    first_value = first_value.title()
                                values.append(first_value)
                    
                    # Pattern 2b: "leading service providers" -> extract values
                    # Pattern: "X, Y, and Z are the leading..."
                    leading_match = re.search(r'^([^.]+?)\s+(?:are|is)\s+the\s+(?:leading|dominant)', header_para, re.IGNORECASE)
                    if leading_match:
                        value_text = leading_match.group(1).strip()
                        value_parts = re.split(r',\s*(?:and\s+)?', value_text)
                        for part in value_parts:
                            part = part.strip()
                            if part and len(part) > 3 and part[0].isupper():
                                if len(part) < 80:
                                    values.append(part)
                    
                    # Pattern 2c: "finds usage in X, Y, Z" or "finds usage in X, Y, and Z"
                    usage_match = re.search(r'finds usage in\s+([^.]*?)(?:\.|$)', header_para, re.IGNORECASE)
                    if usage_match:
                        value_text = usage_match.group(1).strip()
                        # Extract values, handling parentheses like "(ARDS)"
                        # Split by comma, but handle "and" before last item
                        value_parts = re.split(r',\s*(?:and\s+)?', value_text)
                        for part in value_parts:
                            part = part.strip()
                            if part:
                                # Extract main value (may include parentheses for abbreviations)
                                # Keep the full value including parentheses for abbreviations
                                if '(' in part and ')' in part:
                                    # Extract the part before parentheses
                                    main_value = part.split('(')[0].strip()
                                    # Check if it's a valid abbreviation pattern
                                    abbrev_match = re.search(r'\(([A-Z]+)\)', part)
                                    if abbrev_match and len(abbrev_match.group(1)) >= 3:
                                        # Keep with abbreviation
                                        main_value = part
                                else:
                                    main_value = part
                                # Remove trailing description
                                main_value = re.split(r'\s+(?:is|are|projected|expected|account|driven|given)', main_value, 1)[0].strip()
                                if main_value and len(main_value) > 3:
                                    # Capitalize first letter if lowercase
                                    if main_value[0].islower():
                                        # Title case for multi-word phrases
                                        main_value = main_value.title()
                                    # If it has parentheses, keep the abbreviation part
                                    if '(' in main_value and ')' in main_value:
                                        # Extract abbreviation and capitalize the main part
                                        abbrev_match = re.search(r'\(([A-Za-z]+)\)', main_value)
                                        if abbrev_match:
                                            main_part = main_value.split('(')[0].strip()
                                            abbrev = abbrev_match.group(1).upper()  # Uppercase abbreviation
                                            # Title case the main part, but handle multi-word properly
                                            main_part = main_part.title()
                                            # Fix common issues: "And" -> "and", "To" -> "to", "Of" -> "of"
                                            # But keep "To" in phrases like "Bridge-to-Lung" as lowercase
                                            main_part = re.sub(r'\b(And|To|Of|In|For|With|The)\b', lambda m: m.group(1).lower(), main_part)
                                            # Fix hyphenated phrases: "Bridge-To-Lung" -> "Bridge-to-Lung" or "Trauma-Induced" -> "Trauma-Induced" (keep "Induced" capitalized)
                                            # Handle "Bridge-To-Lung" pattern
                                            main_part = re.sub(r'([A-Z][a-z]+)-To-([A-Z][a-z]+)', r'\1-to-\2', main_part)
                                            # Handle "Bridge-to-Lung" already correct
                                            main_part = re.sub(r'([A-Z][a-z]+)-to-([A-Z][a-z]+)', r'\1-to-\2', main_part)
                                            # But keep first word capitalized
                                            if main_part:
                                                main_part = main_part[0].upper() + main_part[1:]
                                            main_value = f"{main_part} ({abbrev})"
                                    if len(main_value) < 100:  # Increased limit for longer values
                                        values.append(main_value)

                    # Pattern 2d: "find applications across X, Y, and Z" / "applications across X, Y, Z"
                    apps_match = re.search(r'(?:find(?:s)?\s+)?applications\s+across\s+([^.]*?)(?:\.|$)', header_para, re.IGNORECASE)
                    if apps_match:
                        value_text = apps_match.group(1).strip()
                        value_text = re.sub(r'^(?:a\s+wide\s+range\s+of\s+)?', '', value_text, flags=re.I).strip()
                        value_parts = re.split(r',\s*(?:and\s+)?', value_text)
                        for part in value_parts:
                            part = part.strip()
                            if not part:
                                continue
                            main_value = re.split(r'\s+(?:lead|dominat|remain|represent|account|continue|expand)', part, 1, flags=re.I)[0].strip()
                            if main_value and len(main_value) < 100:
                                # Title-case common nouns but keep acronyms/ampersands
                                if main_value[0].islower():
                                    main_value = main_value.title()
                                values.append(main_value)

                    # Pattern 2e: "available across a wide spectrum – X, Y, and Z"
                    avail_match = re.search(
                        r'available\s+across\s+(?:a\s+wide\s+spectrum\s*)?[–\-]\s*([^.]*?)(?:\.|$)',
                        header_para,
                        re.IGNORECASE,
                    )
                    if not avail_match:
                        avail_match = re.search(r'available\s+across\s+([^.]*?)(?:\.|$)', header_para, re.IGNORECASE)
                    if avail_match:
                        value_text = avail_match.group(1).strip()
                        # Strip leading filler like "a wide spectrum –"
                        value_text = re.sub(r'^(?:a\s+wide\s+spectrum\s*)?[–\-]\s*', '', value_text, flags=re.I).strip()
                        value_parts = re.split(r',\s*(?:and\s+)?', value_text)
                        for part in value_parts:
                            part = part.strip()
                            if not part:
                                continue
                            main_value = re.split(r'\s+(?:serve|serves|remain|represent|account|are|is)\b', part, 1, flags=re.I)[0].strip()
                            if main_value and len(main_value) < 80:
                                values.append(main_value)
                    
                    # Pattern 3: "spans X, Y, Z" or "spans X, Y, and Z" or "adoption spans X, Y, Z"
                    spans_match = re.search(r'(?:adoption\s+)?spans\s+([^.]*?)(?:\.|$)', header_para, re.IGNORECASE)
                    if spans_match:
                        value_text = spans_match.group(1).strip()
                        value_parts = re.split(r',\s*(?:and\s+)?', value_text)
                        for part in value_parts:
                            part = part.strip()
                            if part:
                                # Remove trailing description
                                main_value = re.split(r'\s+(?:given|dominate|show|large|tertiary)', part, 1)[0].strip()
                                if main_value and len(main_value) > 3:
                                    # Capitalize first letter if lowercase - title case for multi-word
                                    if main_value[0].islower():
                                        main_value = main_value.title()
                                        # Fix common issues: "And" -> "and", "To" -> "to", "Of" -> "of"
                                        main_value = re.sub(r'\b(And|To|Of|In|For|With|The)\b', lambda m: m.group(1).lower(), main_value)
                                        # Fix hyphenated phrases: "Bridge-To-Lung" -> "Bridge-to-Lung"
                                        main_value = re.sub(r'([A-Z][a-z]+)-To-([A-Z][a-z]+)', r'\1-to-\2', main_value)
                                        main_value = re.sub(r'([A-Z][a-z]+)-to-([A-Z][a-z]+)', r'\1-to-\2', main_value)
                                        # But keep first word capitalized
                                        if main_value:
                                            main_value = main_value[0].upper() + main_value[1:]
                                    if len(main_value) < 100:  # Increased limit
                                        values.append(main_value)
                    
                    # Pattern 4a: For region segment - "North America represents the largest market..."
                    if section_type == 'region' and 'represents' in header_lower:
                        # Extract first capitalized phrase (region name)
                        region_match = re.search(r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)', header_para)
                        if region_match:
                            region_value = region_match.group(1).strip()
                            # Remove trailing words before "represents"
                            region_value = re.split(r'\s+(?:represents|is|remains|follows|supported)', region_value, 1)[0].strip()
                            if region_value and len(region_value) < 50 and region_value.lower() not in ['the', 'by', 'market']:
                                # Title case
                                if region_value[0].islower():
                                    region_value = region_value.title()
                                if region_value not in values:
                                    values.append(region_value)
                    
                    # Pattern 4: "distributed across X, Y, Z"
                    distributed_match = re.search(r'distributed across\s+([^.]*?)(?:\.|$)', header_para, re.IGNORECASE)
                    if distributed_match:
                        value_text = distributed_match.group(1).strip()
                        # Handle "X, Y, and Z (LAMEA)" pattern
                        value_parts = re.split(r',\s*(?:and\s+)?(?!\s*\([A-Z]+\))', value_text)
                        for part in value_parts:
                            part = part.strip()
                            if part:
                                # Remove parentheses like "(LAMEA)"
                                main_value = re.sub(r'\s*\([^)]+\)', '', part).strip()
                                if main_value and len(main_value) > 3:
                                    # Capitalize first letter if lowercase
                                    if main_value[0].islower():
                                        main_value = main_value[0].upper() + main_value[1:]
                                    if len(main_value) < 80:
                                        values.append(main_value)
                    
                    # If we confidently extracted a list from the header, don't keep mining narrative
                    # paragraphs (avoids picking sentences like "The X segment is expanding quickly").
                    if values and section_type in {'treatment', 'phase', 'output_power'} and len(values) >= 2:
                        # De-dupe while preserving order
                        deduped = []
                        seen = set()
                        for v in values:
                            vv = v.strip()
                            if not vv:
                                continue
                            if vv.lower() in seen:
                                continue
                            seen.add(vv.lower())
                            deduped.append(vv)
                        return deduped[:8]

                    # If we found values in header, use them and continue
                    if values:
                        # Continue to check following paragraphs for additional values
                        pass
            
            # Also check if header paragraph has values in the same paragraph (after newline)
            # Pattern: "By Product Type\nValue1, Value2, Value3" or "By Product Type\nValue1\nValue2"
            if header_para and '\n' in header_para:
                header_lines = header_para.split('\n')
                if len(header_lines) > 1:
                    # Check lines after the header
                    for line in header_lines[1:]:
                        line = line.strip()
                        if line and len(line) > 10:
                            # Look for comma-separated values
                            if ',' in line:
                                # Extract values from comma-separated list
                                line_values = re.split(r',\s*(?:and\s+)?', line)
                                for val in line_values:
                                    val = val.strip()
                                    # Extract first part before description words
                                    val = re.split(r'\s+(?:remain|are|encompass|represent|account|drive|continue|include|comprise)', val, 1, flags=re.I)[0].strip()
                                    if val and len(val) > 3 and val[0].isupper():
                                        if len(val) < 80:
                                            # Capitalize properly
                                            if val[0].islower():
                                                val = val[0].upper() + val[1:]
                                            if val not in values:
                                                values.append(val)
            
            for i in range(para_idx + 1, min(para_idx + max_paras + 1, len(doc.paragraphs))):
                text = doc.paragraphs[i].text.strip()
                if not text:
                    continue
                
                # Check if this paragraph starts a new section (starts with "By ")
                if re.match(r'^By\s+', text, re.I):
                    break
                
                # Skip very short text (less than 3 chars) unless it's a single capitalized word
                if len(text) < 3:
                    continue
                
                # Handle paragraphs that might contain newlines - extract first line as potential heading
                # Split by newline and check first line
                first_line = text.split('\n')[0].strip() if '\n' in text else text
                
                # Also check if first line contains comma-separated values (might be list of segment values)
                if ',' in first_line and len(first_line) < 200:
                    # Try to extract values from comma-separated list at start of paragraph
                    # Pattern: "Value1, Value2, and Value3 represent..." or "Value1, Value2, Value3"
                    value_match = re.match(r'^([A-Z][^,]+(?:,\s*[A-Z][^,]+){0,5}(?:,\s*and\s+[A-Z][^,]+)?)', first_line)
                    if value_match:
                        value_text = value_match.group(1).strip()
                        # Remove trailing "and" if present
                        value_text = re.sub(r',\s*and\s+([A-Z][^,]+)', r', \1', value_text)
                        value_parts = re.split(r',\s*', value_text)
                        for val in value_parts:
                            val = val.strip()
                            # Extract before description words
                            val = re.split(r'\s+(?:remain|are|encompass|represent|account|drive|continue|include|comprise)', val, 1, flags=re.I)[0].strip()
                            if val and len(val) > 3 and val[0].isupper():
                                if len(val) < 80:
                                    if val not in values:
                                        values.append(val)
                
                text = first_line  # Use first line for extraction
                
                text_lower = text.lower()
                
                # Skip introductory paragraphs (usually longer, don't start with capital word + colon)
                # But be more lenient - only skip if it's clearly an intro paragraph
                # Check if it starts with a capitalized heading first - if so, don't skip
                starts_with_capital_heading = re.match(r'^[A-Z][a-zA-Z0-9\s&\-]+(?:\([A-Za-z0-9\s,\-]+\))?', text)
                if skip_intro and len(text) > 100 and not starts_with_capital_heading:
                    skip_intro = False
                    continue
                skip_intro = False
                
                # Skip if starts with excluded phrases
                if any(text_lower.startswith(phrase) for phrase in exclude_phrases):
                    continue
                
                # Look for capitalized words that might be segment values
                # Pattern: "Technology Name:" or "Technology Name (Description)" - must start with capital and end with colon
                match = re.match(r'^([A-Z][a-zA-Z0-9\s&]+(?:\([A-Za-z0-9\s,]+\))?)\s*[:]', text)
                if match:
                    value = match.group(1).strip()
                    # Remove parentheses content for cleaner output (e.g., "Canned Fish (Tuna, Salmon, Sardines)" -> "Canned Fish")
                    # But keep it if it's an abbreviation like "(PGPR)" or "(AML)"
                    if '(' in value and ')' in value:
                        paren_match = re.search(r'\(([^)]+)\)', value)
                        if paren_match:
                            paren_content = paren_match.group(1).strip()
                            # If it's all uppercase letters (abbreviation), keep it
                            # Otherwise, remove the parentheses content
                            if not re.match(r'^[A-Z]{2,}$', paren_content):
                                value = re.sub(r'\s*\([^)]+\)', '', value).strip()
                    # Clean up value - remove common prefixes
                    value = re.sub(r'^(The|This|These|That)\s+', '', value, flags=re.I).strip()
                    value_lower = value.lower()
                    
                    # Check if it's a region name
                    is_region_name = any(region in value_lower for region in region_names)
                    
                    # Skip region names if we're not extracting for the region section
                    if is_region_name and section_type != 'region':
                        continue
                    
                    # For region section, only accept known region names or region-like patterns
                    if section_type == 'region':
                        # Must be a known region or look like a region name
                        if not is_region_name:
                            # Check if it looks like a region name (short, capitalized, no verbs)
                            if (len(value) > 40 or 
                                any(verb in value_lower for verb in ['influence', 'market size', 'volume', 'diagnostics', 'methodology', 'process', 'research', 'data sources'])):
                                continue  # Skip non-region values
                    
                    # Check if it's a valid value - more flexible for generic markets
                    # Allow longer values (up to 60 chars) and more words (up to 8) for compound names
                    is_valid = (len(value) < 70 and 
                               not value_lower.startswith(('by ', 'the ', 'this ', 'these ', 'aml ', 'market', 'key ', 'the global', 'other ', 'projected ')) and
                               len(value.split()) <= 10 and  # Increased for compound names like "Food & Beverages"
                               value not in values and
                               'market analysis' not in value_lower and
                               'market share' not in value_lower and
                               'projected share' not in value_lower and
                               'the market' not in value_lower and
                               'the end users' not in value_lower and
                               'end-user' not in value_lower and
                               not value_lower.startswith('oem') and
                               not value_lower.startswith('rest of'))  # Exclude "Rest of the World"
                    
                    # Allow "Diagnostic Laboratories" and similar compound terms, but exclude standalone "Diagnostic"
                    if is_valid:
                        if value_lower.startswith('diagnostic'):
                            # Allow if it's a compound term like "Diagnostic Laboratories"
                            if 'laborator' in value_lower or 'service' in value_lower or 'center' in value_lower:
                                values.append(value)
                        else:
                            # Normalize "and" to "&" for compound names if needed
                            # But keep "and" for regions like "Middle East & Africa"
                            if ' and ' in value and section_type != 'region':
                                # For application/industry segments, use &
                                value = value.replace(' and ', ' & ')
                            values.append(value)
                
                # Also check for values that might appear at the start of longer paragraphs
                # Pattern: First few words before colon, even if followed by more text
                # This handles cases like "Food and Beverages: The largest application segment..."
                if not match and len(text) > 50:  # Only for longer paragraphs
                    # Look for pattern: Capital words followed by colon, within first 100 chars
                    text_start = text[:100]
                    alt_match = re.search(r'^([A-Z][a-zA-Z0-9\s&]+(?:\([A-Za-z0-9\s,]+\))?)\s*[:]', text_start)
                    if alt_match:
                        value = alt_match.group(1).strip()
                        
                        # Skip if it looks like a sentence with verbs (description text)
                        value_lower_check = value.lower()
                        sentence_indicators = ['is employed', 'are used', 'is used', 'are employed', 'focuses on', 'targeting', 'including', 'such as']
                        if any(indicator in text_lower for indicator in sentence_indicators):
                            # This is likely a description, not a segment value
                            pass
                        else:
                            # Remove parentheses content for cleaner output (e.g., "Canned Fish (Tuna, Salmon, Sardines)" -> "Canned Fish")
                            # But keep it if it's an abbreviation like "(PGPR)" or "(AML)"
                            if '(' in value and ')' in value:
                                paren_match = re.search(r'\(([^)]+)\)', value)
                                if paren_match:
                                    paren_content = paren_match.group(1).strip()
                                    # If it's all uppercase letters (abbreviation), keep it
                                    # Otherwise, remove the parentheses content
                                    if not re.match(r'^[A-Z]{2,}$', paren_content):
                                        value = re.sub(r'\s*\([^)]+\)', '', value).strip()
                            value = re.sub(r'^(The|This|These|That)\s+', '', value, flags=re.I).strip()
                            value_lower = value.lower()
                            
                            # Same validation as above
                            is_region_name = any(region in value_lower for region in region_names)
                            if is_region_name and section_type != 'region':
                                pass  # Skip
                            elif (len(value) < 70 and 
                                  not value_lower.startswith(('by ', 'the ', 'this ', 'these ', 'market', 'key ', 'other ', 'rest of', 'projected ')) and
                                  len(value.split()) <= 10 and
                                  value not in values and
                                  'market analysis' not in value_lower and
                                  'market share' not in value_lower and
                                  'projected share' not in value_lower):
                                # Normalize "and" to "&" for compound names
                                if ' and ' in value and section_type != 'region':
                                    value = value.replace(' and ', ' & ')
                                values.append(value)
                
                # Also check for standalone capitalized headings (without colons)
                # Pattern: Text that starts with capital letter, is reasonably short, and looks like a heading
                # This handles cases where headings appear on their own line or followed by description on same line
                if not match and len(text) > 0:
                    # Check if text starts with a capitalized heading pattern
                    # Pattern: Capital words that might be a heading, possibly followed by description
                    # Extract just the first capitalized phrase (the heading part) - be more flexible
                    # Try to match up to first sentence break or newline
                    heading_match = None
                    # First try: exact match for standalone heading (no description)
                    if re.match(r'^[A-Z][a-zA-Z0-9\s&\-]+(?:\([A-Za-z0-9\s,\-]+\))?$', text.strip()):
                        heading_match = text.strip()
                    else:
                        # Second try: heading followed by description - extract first part
                        # More flexible pattern: match capitalized words until we hit lowercase word (description starts)
                        # Pattern: Match capitalized words, hyphens, and spaces until we hit a lowercase letter
                        # This handles "Bone Marrow Failure Syndromes\nPatients..." or "Post-Hematopoietic Stem Cell Transplantation\nG-CSF..."
                        heading_pattern = re.match(r'^([A-Z][A-Za-z0-9\s&\-]+(?:\([A-Za-z0-9\s,\-]+\))?)', text)
                        if heading_pattern:
                            potential_heading = heading_pattern.group(1).strip()
                            # Check if what follows is description (starts with lowercase or is end of text)
                            remaining = text[len(potential_heading):].strip()
                            if not remaining:
                                # No description, this is the full heading
                                heading_match = potential_heading
                            elif remaining and remaining[0].islower():
                                # Description starts with lowercase, heading is complete
                                heading_match = potential_heading
                            elif len(potential_heading.split()) <= 8 and len(potential_heading) <= 80:
                                # Heading seems reasonable in length
                                heading_match = potential_heading
                    
                    if heading_match:
                        potential_value = heading_match
                        value_lower = potential_value.lower()
                        
                        # Check if this looks like a standalone heading value
                        # Must be reasonably short (not a full paragraph)
                        is_heading_like = (
                            len(potential_value) <= 80 and  # Not too long
                            len(potential_value.split()) <= 8 and  # Not too many words
                            not re.match(r'^\d+[\.\)]\s*', potential_value) and  # Not numbered
                            not value_lower.startswith(('by ', 'the ', 'this ', 'these ', 'market', 'key ', 'projected ', 'products such', 'the g-csf market')) and
                            'market analysis' not in value_lower and
                            'market share' not in value_lower and
                            'projected share' not in value_lower and
                            potential_value not in values
                        )
                        
                        if is_heading_like:
                            # Check if it's a region name (skip if not extracting for region)
                            is_region_name = any(region in value_lower for region in region_names)
                            if is_region_name and section_type != 'region':
                                pass  # Skip
                            else:
                                # Clean up the value
                                value = re.sub(r'^(The|This|These|That)\s+', '', potential_value, flags=re.I).strip()
                                value_lower = value.lower()
                                
                                # Skip if it contains sentence verbs/patterns (description text)
                                sentence_indicators = ['is employed', 'are used', 'is used', 'are employed', 'focuses on', 
                                                      'targeting', 'including', 'such as', 'helps in', 'provides', 
                                                      'across various', 'each targeting', 'each with']
                                if any(indicator in value_lower for indicator in sentence_indicators):
                                    pass  # Skip description text
                                # Skip if it contains phrases that indicate it's not a segment value
                                elif any(phrase in value_lower for phrase in ['increased', 'access to', 'growth rate', 'expected to', 'impact', 'trend', 'outlook', 'driven by', 'the largest', 'the dominant', 'patients with', 'g-csf is', 'targeted', 'treatments', 'improve patient', 'outcomes', 'healthcare', 'infrastructure']):
                                    pass  # Skip descriptive phrases
                                # Skip if it's a single generic word that's not a proper segment value
                                elif len(value.split()) == 1 and value_lower in ['impact', 'trend', 'outlook', 'growth', 'share', 'market', 'targeted']:
                                    pass  # Skip generic single words
                                # Skip if it's too long and looks like a sentence/description
                                elif len(value) > 60:
                                    # Check for sentence structure indicators
                                    if any(word in value_lower for word in ['to', 'for', 'with']) and ('across' in value_lower or 'various' in value_lower or 'each' in value_lower):
                                        pass  # Likely a descriptive phrase, not a segment value
                                    elif len(value.split()) > 8:  # Too many words
                                        pass
                                    elif value[0].islower():  # Starts with lowercase (likely part of sentence)
                                        pass
                                    elif any(phrase in value_lower for phrase in ['the primary', 'this segment', 'this application']):
                                        pass
                                    else:
                                        # Might be valid, but validate further
                                        if len(value) >= 3 and len(value) <= 80:
                                            values.append(value)
                                # Only add if it's a reasonable heading (not too generic)
                                elif len(value) >= 3 and len(value) <= 80:
                                    values.append(value)
            
            # For region section, handle "Rest of the World" or "Latin America, Middle East & Africa" specially
            if section_type == 'region' and values:
                # Filter to only keep valid region names
                valid_regions = []
                region_keywords = ['north america', 'south america', 'europe', 'asia', 'pacific', 'latin america', 
                                 'middle east', 'africa', 'lamea', 'apac', 'emea', 'america', 'oceania']
                
                for val in values:
                    val_lower = val.lower()
                    # Check if it's a known region or contains region keywords
                    if any(region in val_lower for region in region_names + region_keywords):
                        # Additional validation: exclude non-region terms
                        invalid_region_terms = ['influence', 'market size', 'volume', 'diagnostics', 'methodology', 
                                               'process', 'research', 'data sources', 'government', 'regulatory',
                                               'historical', 'overview', 'emerging', 'technological', 'stakeholders']
                        if not any(term in val_lower for term in invalid_region_terms):
                            valid_regions.append(val)
                
                # If we found valid regions, use them; otherwise keep original values
                if valid_regions:
                    values = valid_regions
                
                # Check if we have "Rest of the World" pattern
                for i, val in enumerate(values):
                    if 'rest of' in val.lower() or 'lamea' in val.lower():
                        # Look for expanded form in document
                        for para in doc.paragraphs[para_idx:para_idx+max_paras+5]:
                            para_text = para.text.lower()
                            if 'latin america' in para_text and 'middle east' in para_text and 'africa' in para_text:
                                # Replace with proper format
                                values[i] = 'Latin America, Middle East & Africa'
                                break
            
            # Clean up all values - remove any invalid ones
            final_values = []
            # Common verbs and sentence patterns that indicate description text, not segment values
            sentence_verbs = ['is employed', 'are used', 'is used', 'are employed', 'focuses on', 'targeting', 
                            'including', 'such as', 'helps in', 'provides', 'allows', 'enables', 'ensures',
                            'aims to', 'seeks to', 'designed to', 'used to', 'intended to']
            
            # Region-specific validation
            region_keywords = ['north america', 'south america', 'europe', 'asia', 'pacific', 'latin america', 
                             'middle east', 'africa', 'lamea', 'apac', 'emea', 'america', 'oceania']
            invalid_region_terms = ['influence', 'market size', 'volume', 'diagnostics', 'methodology', 
                                   'process', 'research', 'data sources', 'government', 'regulatory',
                                   'historical', 'overview', 'emerging', 'technological', 'stakeholders']
            
            for val in values:
                # Strip "(e.g., X)" from segment values so we get short form (e.g. "Internal Stabilization Systems" not "Internal Stabilization Systems (e.g., Internal...")
                if "(e.g." in val or "(e.g.," in val:
                    val = val.split("(e.g.")[0].strip().split("(e.g.,")[0].strip()
                if not val:
                    continue
                val_lower = val.lower()
                
                # Special handling for region section
                if section_type == 'region':
                    # Must contain region keywords
                    if not any(region in val_lower for region in region_keywords):
                        continue
                    # Must not contain invalid terms
                    if any(term in val_lower for term in invalid_region_terms):
                        continue
                
                # Skip if it contains sentence verbs/patterns (description text)
                if any(verb in val_lower for verb in sentence_verbs):
                    continue
                    
                # Skip if it's clearly not a valid segment value
                invalid_patterns = [
                    'impact', 'trend', 'outlook', 'growth', 'share', 'targeted',
                    'increased', 'access to', 'growth rate', 'expected to', 'driven by',
                    'the largest', 'improve patient', 'treatments to', 'outcomes',
                    'healthcare infrastructure', 'reflects', 'differences in',
                    'definition and scope', 'overview of', 'research methodology', 'research process',
                    'primary and secondary', 'data sources', 'emerging opportunities', 'technological advances',
                    'across various', 'each targeting', 'each with', 'vary widely',
                    'government and regulatory', 'historical market', 'market size and volume',
                    'this is the most', 'this is the critical', 'dimension of segmentation',
                    'in 2024,', 'in 2024 ', 'the most critical dimension',
                ]
                if (val_lower in invalid_patterns[:6] or 
                    any(phrase in val_lower for phrase in invalid_patterns[6:])):
                    continue
                    
                # Skip if it's too long and looks like a description/sentence
                if len(val) > 60:
                    # Check for sentence structure (contains verbs, prepositions indicating description)
                    if any(word in val_lower for word in ['to', 'for', 'with', 'and', 'of', 'in', 'at', 'on', 'by']):
                        # Count prepositions - if too many, it's likely a sentence
                        prep_count = sum(1 for word in ['to', 'for', 'with', 'of', 'in', 'at', 'on', 'by'] if word in val_lower.split())
                        if prep_count > 2:
                            continue
                
                # Skip if it starts with lowercase (likely part of a sentence)
                if val and val[0].islower():
                    continue
                    
                # Skip if it contains common description phrases
                desc_phrases = ['the primary', 'this segment', 'this sub-segment', 'this application',
                              'the end users', 'the market', 'the largest', 'the dominant',
                              'this is the', 'the most critical']
                if any(phrase in val_lower for phrase in desc_phrases):
                    continue
                
                # Allow valid values
                final_values.append(val)
            
            return final_values[:8]  # Increased limit to 8 values for comprehensive segments
        
        # Construct base market name
        # Don't replace hyphens inside parentheses (e.g., "(G-CSF)" should stay as is)
        # First, preserve abbreviations in parentheses by splitting and processing separately
        preserved_abbrev = None  # Initialize to avoid UnboundLocalError
        abbrev_pattern_file = re.search(r'(\([A-Z0-9\-]+\))', filename)
        if abbrev_pattern_file:
            preserved_abbrev = abbrev_pattern_file.group(1)
            # Split filename into parts before and after abbreviation
            parts = filename.split(preserved_abbrev, 1)
            # Process each part separately (replace hyphens/underscores)
            before = parts[0].replace('-', ' ').replace('_', ' ').strip()
            after = parts[1].replace('-', ' ').replace('_', ' ').strip() if len(parts) > 1 else ''
            # Reconstruct with preserved abbreviation
            base_market_name = f"{before} {preserved_abbrev} {after}".strip()
        else:
            base_market_name = filename.replace('-', ' ').replace('_', ' ')
        
        # Add "Global" prefix if not already present
        # if not base_market_name.lower().startswith('global'):
        #     base_market_name = f"Global {base_market_name}"
        
        # Check if document mentions abbreviation in first few paragraphs
        # Look for pattern like "Polyglycerol Polyricinoleate (PGPR)" or "Acute Myeloid Leukemia (AML)"
        abbreviation = None
        filename_lower = base_market_name.lower()
        
        # Extract market name keywords for abbreviation matching
        market_keywords = [w for w in filename_lower.replace(' market', '').split() if len(w) > 3]
        
        for i, para in enumerate(doc.paragraphs[:20]):  # Check more paragraphs
            text = para.text
            text_lower = text.lower()
            
            # Check if paragraph contains market name keywords
            has_market_keywords = any(kw in text_lower for kw in market_keywords)
            
            if has_market_keywords:
                # Look for abbreviation pattern: "Market Name (ABBR)" or "Market Name(ABBR)"
                # Pattern: word(s) followed by (uppercase letters, may include hyphens like "G-CSF")
                abbrev_pattern = re.search(r'\(([A-Z][A-Z0-9\-]{1,})\)', text)
                if abbrev_pattern:
                    abbrev = abbrev_pattern.group(1)
                    # Verify it's near market name
                    abbrev_pos = abbrev_pattern.start()
                    market_context = text[max(0, abbrev_pos-100):abbrev_pos+50].lower()
                    if any(kw in market_context for kw in market_keywords):
                        abbreviation = f'({abbrev})'
                        break
                
                # Also check for specific known abbreviations
                if 'polyglycerol polyricinoleate' in text_lower or 'pgpr' in text_lower:
                    pgpr_match = re.search(r'\(PGPR\)', text, re.I)
                    if pgpr_match:
                        abbreviation = '(PGPR)'
                        break
                elif 'acute myeloid leukemia' in text_lower or 'aml diagnostics' in text_lower:
                    aml_match = re.search(r'\(AML\)', text, re.I)
                    if aml_match:
                        abbreviation = '(AML)'
                        break
        
        # If abbreviation is already in the name (from filename), don't add it again
        # Also check if abbreviation is different from what's already in filename
        if preserved_abbrev and abbreviation:
            # Check if they match (case-insensitive)
            if preserved_abbrev.lower() == abbreviation.lower():
                abbreviation = None  # Don't add duplicate
            # Also check if abbreviation is already mentioned in the base_market_name
            elif abbreviation.lower() in base_market_name.lower():
                abbreviation = None  # Don't add if already present
        
        # Add abbreviation after disease name but before other words if found
        # Check if abbreviation is already in the base_market_name (from filename)
        abbrev_in_name = abbreviation and abbreviation.lower() in base_market_name.lower()
        # Also check if base_market_name already has an abbreviation (from filename) - don't add duplicate
        if preserved_abbrev:
            # If filename already has an abbreviation (like "(VV ECLS)"), don't add another one (like "(ECMO)")
            # Only exception: G-CSF format where we need to rearrange
            abbrev_in_name = True
            # Reset abbreviation to None to prevent adding duplicate
            abbreviation = None
        elif abbreviation and abbreviation.lower() in base_market_name.lower():
            # If abbreviation is already in base_market_name, don't add it
            abbrev_in_name = True
            abbreviation = None
        if abbreviation and not abbrev_in_name:
            # Insert after the disease name (before "Diagnostics" or "Market")
            # Pattern: "Acute Myeloid Leukemia (AML) Diagnostics Market"
            if 'diagnostics' in base_market_name.lower():
                base_market_name = base_market_name.replace(' Diagnostics', f' {abbreviation} Diagnostics')
            elif 'market' in base_market_name.lower():
                base_market_name = base_market_name.replace(' Market', f' {abbreviation} Market')
            else:
                base_market_name = f"{base_market_name} {abbreviation}"
        
        if not base_market_name.lower().endswith('market'):
            base_market_name = f"{base_market_name} Market"
        
        # Special handling for G-CSF format: "G-CSF (Granulocyte Colony Stimulating Factors) Market"
        # Check if we have abbreviation and full name pattern
        if preserved_abbrev and '(G-CSF)' in preserved_abbrev.upper():
            # Check if full name is in the base_market_name
            if 'granulocyte colony stimulating factors' in base_market_name.lower():
                # Extract abbreviation and full name
                abbrev = 'G-CSF'
                full_name = 'Granulocyte Colony Stimulating Factors'
                # Reconstruct as "G-CSF (Granulocyte Colony Stimulating Factors) Market"
                base_market_name = f"{abbrev} ({full_name}) Market"
        
        # Build the detailed title
        title_parts = [base_market_name]
        
        # Build title parts in specific order: Phase Type, Output Power, Product Type, Distribution Channel, Application, End-User, Geography
        # Note: Order matters for consistency

        # 0. Phase Type (for power/electrical markets)
        if 'phase' in segments and 'phase' in segment_indices:
            phase_values = extract_segment_values(segment_indices['phase'], 'phase')
            if phase_values:
                cleaned_values = []
                for val in phase_values:
                    v = val.strip()
                    if not v:
                        continue
                    # normalize capitalization for common values
                    if v.lower() == 'single phase':
                        v = 'Single Phase'
                    elif v.lower() == 'three phase':
                        v = 'Three Phase'
                    if v not in cleaned_values:
                        cleaned_values.append(v)
                if cleaned_values:
                    title_parts.append(f"By Phase Type ({', '.join(cleaned_values[:6])})")

        # 0.5 Output Power (for power/electrical markets)
        if 'output_power' in segments and 'output_power' in segment_indices:
            output_values = extract_segment_values(segment_indices['output_power'], 'output_power')
            if output_values:
                cleaned_values = []
                for val in output_values:
                    v = _norm(val)
                    if not v:
                        continue
                    # ensure we use en-dash between numeric ranges
                    v = re.sub(r'(\d)\s*[-]\s*(\d)', rf'\1{DASH}\2', v)
                    if v not in cleaned_values:
                        cleaned_values.append(v)
                if cleaned_values:
                    title_parts.append(f"By Output Power ({', '.join(cleaned_values[:6])})")
        
        # 1. Technology/Type (Diagnostic Technology or generic Product Type)
        if 'diagnostic' in segments:
            if 'diagnostic' in segment_indices:
                diagnostic_values = extract_segment_values(segment_indices['diagnostic'], 'diagnostic')
                if diagnostic_values:
                    # Check if this is medical-specific or generic segmentation
                    segment_text = segments['diagnostic'].lower()
                    is_medical = 'diagnostic' in segment_text and ('technology' in segment_text or 'approach' in segment_text)
                    is_generic_type = 'type' in segment_text or 'product type' in segment_text
                    
                    # Clean up values - map to expected names for medical, preserve for generic
                    cleaned_values = []
                    
                    if is_medical:
                        # Medical-specific mapping
                        expected_order = ['Molecular Diagnostics', 'Flow Cytometry', 'NGS', 'Liquid Biopsy', 'IHC', 'Others']
                    for val in diagnostic_values:
                        val_lower = val.lower()
                        if 'molecular diagnostics' in val_lower and 'Molecular Diagnostics' not in cleaned_values:
                            cleaned_values.append('Molecular Diagnostics')
                        if 'flow cytometry' in val_lower and 'Flow Cytometry' not in cleaned_values:
                            cleaned_values.append('Flow Cytometry')
                        if ('ngs' in val_lower or 'next generation sequencing' in val_lower or 'next-gen sequencing' in val_lower) and 'NGS' not in cleaned_values:
                            cleaned_values.append('NGS')
                        if 'liquid biopsy' in val_lower and 'Liquid Biopsy' not in cleaned_values:
                            cleaned_values.append('Liquid Biopsy')
                        if ('immunohistochemistry' in val_lower or ('ihc' in val_lower and val_lower != 'others')) and 'IHC' not in cleaned_values:
                            cleaned_values.append('IHC')
                        if ('others' in val_lower or val == 'Others') and 'Others' not in cleaned_values:
                            cleaned_values.append('Others')
                    
                        # Also search document-wide for NGS and Liquid Biopsy if not found
                    if 'NGS' not in cleaned_values:
                        for para in doc.paragraphs:
                            text_lower = para.text.lower()
                            if ('ngs' in text_lower or 'next generation sequencing' in text_lower) and 'diagnostic' in text_lower[:100]:
                                cleaned_values.append('NGS')
                                break
                    if 'Liquid Biopsy' not in cleaned_values:
                        for para in doc.paragraphs:
                            text_lower = para.text.lower()
                            if 'liquid biopsy' in text_lower and 'diagnostic' in text_lower[:100]:
                                cleaned_values.append('Liquid Biopsy')
                                break
                    
                        # Use "By Technology" for medical markets
                    if cleaned_values:
                        ordered_values = []
                        preferred_order = ['Molecular Diagnostics', 'Flow Cytometry', 'NGS', 'Liquid Biopsy', 'IHC', 'Others']
                        for preferred in preferred_order:
                            if preferred in cleaned_values:
                                ordered_values.append(preferred)
                        for val in cleaned_values:
                            if val not in ordered_values:
                                ordered_values.append(val)
                        title_parts.append(f"By Technology ({', '.join(ordered_values[:6])})")
                    elif is_generic_type:
                        # Generic type segmentation - clean up values
                        cleaned_values = []
                        for val in diagnostic_values:
                            val_lower = val.lower()
                            # Map G-CSF specific values
                            if 'innovator' in val_lower and 'g-csf' in val_lower:
                                # Clean "Innovator G-CSF Drugs" to "Innovator G-CSF"
                                if 'drugs' in val_lower:
                                    cleaned_values.append('Innovator G-CSF')
                                else:
                                    cleaned_values.append('Innovator G-CSF')
                            elif 'biosimilar' in val_lower:
                                cleaned_values.append('Biosimilars')
                            elif val not in cleaned_values:
                                cleaned_values.append(val)
                        
                        if cleaned_values:
                            # Use "by Type" for G-CSF markets, "By Product Type" for others
                            segment_text = segments.get('diagnostic', '').lower()
                            if 'type of product' in segment_text or ('g-csf' in base_market_name.lower()):
                                title_parts.append(f"by Type ({', '.join(cleaned_values[:6])})")
                            else:
                                title_parts.append(f"By Product Type ({', '.join(cleaned_values[:6])})")
                    else:
                        # Default: preserve values and use "By Product Type" for generic markets
                        cleaned_values = [val for val in diagnostic_values if len(val) < 50][:6]
                        if cleaned_values:
                            # Use "By Product Type" for generic markets
                            segment_text = segments.get('diagnostic', '').lower()
                            if 'product type' in segment_text or 'type' in segment_text:
                                title_parts.append(f"By Product Type ({', '.join(cleaned_values)})")
                            else:
                                title_parts.append(f"By Product Type ({', '.join(cleaned_values)})")
        
        # 2. Application (Treatment/Application)
        if 'treatment' in segments:
            if 'treatment' in segment_indices:
                treatment_values = extract_segment_values(segment_indices['treatment'], 'treatment')
                if treatment_values:
                    # Clean up values - preserve original for generic markets
                    cleaned_values = []
                    for val in treatment_values:
                        val_lower = val.lower()
                        # Medical-specific mappings
                        if 'disease diagnosis' in val_lower:
                            cleaned_values.append('Disease Diagnosis')
                        elif 'prognostic' in val_lower:
                            cleaned_values.append('Prognostic Determination')
                        elif 'treatment monitoring' in val_lower:
                            cleaned_values.append('Treatment Monitoring')
                        elif 'recurrence' in val_lower:
                            cleaned_values.append('Recurrence Detection')
                        # G-CSF specific mappings
                        elif 'chemotherapy-induced neutropenia' in val_lower:
                            if 'Chemotherapy-Induced Neutropenia' not in cleaned_values:
                                cleaned_values.append('Chemotherapy-Induced Neutropenia')
                        elif 'bone marrow failure' in val_lower:
                            if 'Bone Marrow Failure' not in cleaned_values:
                                cleaned_values.append('Bone Marrow Failure')
                        elif ('stem cell transplantation' in val_lower or 'hematopoietic stem cell' in val_lower or 
                              'post-hematopoietic' in val_lower):
                            if 'Stem Cell Transplantation' not in cleaned_values:
                                cleaned_values.append('Stem Cell Transplantation')
                        elif 'chronic neutropenia' in val_lower:
                            if 'Chronic Neutropenia' not in cleaned_values:
                                cleaned_values.append('Chronic Neutropenia')
                        elif val not in cleaned_values:
                            # For generic markets, preserve as-is but clean up
                            # Fix specific capitalization issues
                            # Fix "Bridge-To-Lung" -> "Bridge-to-Lung"
                            if 'bridge' in val_lower and 'lung' in val_lower:
                                val = re.sub(r'Bridge-To-Lung', 'Bridge-to-Lung', val, flags=re.I)
                            # Map "Other Applications" to "Industrial" if needed
                            if 'other' in val_lower and 'industrial' in val_lower:
                                if 'Industrial' not in cleaned_values:
                                    cleaned_values.append('Industrial')
                            elif 'other application' in val_lower:
                                # Skip "Other Applications" if we have Industrial already
                                if 'Industrial' not in cleaned_values:
                                    cleaned_values.append('Industrial')
                            else:
                                cleaned_values.append(val)
                    if cleaned_values:
                        # Keep consistent capitalization for titles
                        title_parts.append(f"By Application ({', '.join(cleaned_values[:6])})")
        
        # 3. End-User Industry
        if 'enduser' in segments:
            if 'enduser' in segment_indices:
                enduser_values = extract_segment_values(segment_indices['enduser'], 'enduser')
                if enduser_values:
                    # Clean up values - preserve all found values
                    cleaned_values = []
                    value_added = set()
                    
                    for val in enduser_values:
                        val_lower = val.lower()
                        # Medical-specific mappings
                        if 'hospitals' in val_lower:
                            if 'Hospitals' not in cleaned_values:
                                cleaned_values.append('Hospitals')
                        elif 'diagnostic laborator' in val_lower:
                            if 'Diagnostic Laboratories' not in cleaned_values:
                                cleaned_values.append('Diagnostic Laboratories')
                        elif 'research' in val_lower and ('institut' in val_lower or 'academic' in val_lower):
                            if 'Research Institutions' not in cleaned_values:
                                cleaned_values.append('Research Institutions')
                        # G-CSF specific end-user mappings
                        elif 'hospital' in val_lower:
                            if 'Hospitals' not in cleaned_values:
                                cleaned_values.append('Hospitals')
                        elif 'oncology clinic' in val_lower:
                            if 'Oncology Clinics' not in cleaned_values:
                                cleaned_values.append('Oncology Clinics')
                        elif 'ambulatory' in val_lower and ('surgical' in val_lower or 'surgery' in val_lower):
                            if 'Ambulatory Surgical Centers' not in cleaned_values:
                                cleaned_values.append('Ambulatory Surgical Centers')
                        elif 'homecare' in val_lower or 'home care' in val_lower:
                            if 'Homecare Settings' not in cleaned_values:
                                cleaned_values.append('Homecare Settings')
                        elif 'pharmaceutical' in val_lower:
                            if 'Pharmaceutical' not in ' '.join(cleaned_values):
                                # Map to "Pharmaceuticals" or "Pharmaceutical Industry"
                                if 'industry' in val_lower:
                                    cleaned_values.append('Pharmaceuticals')
                                else:
                                    cleaned_values.append('Pharmaceuticals')
                        elif 'food' in val_lower:
                            if 'Food Industry' not in cleaned_values:
                                cleaned_values.append('Food Industry')
                        elif 'cosmetic' in val_lower:
                            if 'Cosmetics & Personal Care' not in cleaned_values:
                                cleaned_values.append('Cosmetics & Personal Care')
                        elif 'agricultural' in val_lower or 'agriculture' in val_lower:
                            if 'Agriculture' not in cleaned_values:
                                cleaned_values.append('Agriculture')
                        elif 'industrial' in val_lower or 'other' in val_lower:
                            if 'Industrial' not in cleaned_values:
                                cleaned_values.append('Industrial')
                        elif val not in cleaned_values:
                            # Preserve original value
                            cleaned_values.append(val)
                    
                    # For G-CSF markets, also search document-wide for "Ambulatory Surgical Centers" if not found
                    if 'g-csf' in base_market_name.lower() and 'Ambulatory Surgical Centers' not in cleaned_values:
                        for para in doc.paragraphs:
                            para_text_lower = para.text.lower()
                            if 'ambulatory' in para_text_lower and ('surgical center' in para_text_lower or 'asc' in para_text_lower):
                                cleaned_values.append('Ambulatory Surgical Centers')
                                break
                    
                    # For G-CSF markets, ensure correct order: Hospitals, Oncology Clinics, Ambulatory Surgical Centers, Homecare Settings
                    if 'g-csf' in base_market_name.lower() and cleaned_values:
                        ordered_enduser = []
                        preferred_order = ['Hospitals', 'Oncology Clinics', 'Ambulatory Surgical Centers', 'Homecare Settings']
                        for preferred in preferred_order:
                            if preferred in cleaned_values:
                                ordered_enduser.append(preferred)
                        # Add any remaining values not in preferred order
                        for val in cleaned_values:
                            if val not in ordered_enduser:
                                ordered_enduser.append(val)
                        cleaned_values = ordered_enduser
                    
                    # Use "by End-User" for G-CSF, "By End User" for others (remove "Industry")
                    if cleaned_values:
                        if 'g-csf' in base_market_name.lower():
                            title_parts.append(f"by End-User ({', '.join(cleaned_values[:6])})")
                        else:
                            # Check if segment text says "By End User" (without Industry)
                            segment_text = segments.get('enduser', '').lower()
                            if 'industry' not in segment_text:
                                title_parts.append(f"By End User ({', '.join(cleaned_values[:6])})")
                            else:
                                title_parts.append(f"By End-User Industry ({', '.join(cleaned_values[:6])})")
        
        # 3.5. Distribution Channel
        if 'distribution' in segments:
            if 'distribution' in segment_indices:
                distribution_values = extract_segment_values(segment_indices['distribution'], 'distribution')
                if distribution_values:
                    # Clean up values
                    cleaned_values = []
                    for val in distribution_values:
                        val_lower = val.lower()
                        # Map to standard names
                        if 'supermarket' in val_lower or 'hypermarket' in val_lower:
                            if 'Supermarkets' not in cleaned_values:
                                cleaned_values.append('Supermarkets')
                        elif 'online' in val_lower and 'retail' in val_lower:
                            if 'Online Retail' not in cleaned_values:
                                cleaned_values.append('Online Retail')
                        elif 'convenience' in val_lower and 'store' in val_lower:
                            if 'Convenience Stores' not in cleaned_values:
                                cleaned_values.append('Convenience Stores')
                        elif 'foodservice' in val_lower or 'food service' in val_lower:
                            if 'Foodservice' not in cleaned_values:
                                cleaned_values.append('Foodservice')
                        elif val not in cleaned_values:
                            cleaned_values.append(val)
                    
                    if cleaned_values:
                        if 'g-csf' in base_market_name.lower():
                            title_parts.append(f"by Distribution Channel ({', '.join(cleaned_values[:6])})")
                        else:
                            title_parts.append(f"By Distribution Channel ({', '.join(cleaned_values[:6])})")
        
        # 4. Region (Geography) - always add if segments found
        # Check if region segment was found in the document
        if 'region' in segments:
            # In long-form report titles, this is typically expressed as "By Geography"
            # (even if the body text lists regions). Prefer this stable label.
            if 'g-csf' in base_market_name.lower():
                title_parts.append("by Region")
            else:
                title_parts.append("By Geography")
        elif len(segments) >= 2:
            # If we have other segments, add Region anyway
            if 'g-csf' in base_market_name.lower():
                title_parts.append("by Region")
            else:
                title_parts.append("By Geography")
        
        # Add standard ending
        title_parts.append("Segment Revenue Estimation, Forecast, 2024–2030")
        
        # Join title parts with special handling: first segment after market name should have space
        # For G-CSF markets, use commas; for others, use semicolons
        is_gcsf_market = 'g-csf' in base_market_name.lower()
        separator = ', ' if is_gcsf_market else '; '
        
        if len(title_parts) > 1:
            # First part is market name, rest are segments
            market_name = title_parts[0]
            segments = title_parts[1:]
            
            # Join: Market Name + space + first segment, then separator for rest
            if segments:
                # First segment should be joined with space after market name
                constructed_title = f"{market_name} {segments[0]}"
                # Rest segments joined with separator
                if len(segments) > 1:
                    constructed_title += separator + separator.join(segments[1:])
            else:
                constructed_title = market_name
        else:
            constructed_title = separator.join(title_parts)
        
        return _clean_final_title(_norm(constructed_title))

    # NEW LOGIC: Look for market report patterns in first few paragraphs
    # Increased to 50 paragraphs to catch titles that appear after intro text
    filename_normalized = filename_low.replace('-', ' ').replace('_', ' ')
    filename_keywords = [w for w in filename_normalized.split() if w not in ['market', 'the', 'and', 'or', 'global']]
    
    # For 2-word market names, require both words to match (more flexible)
    # For longer names, require at least 2 keywords or 70% match
    min_keywords_needed = max(2, min(len(filename_keywords), 3)) if len(filename_keywords) > 2 else len(filename_keywords)
    
    for para_idx, para in enumerate(doc.paragraphs[:50]):  # Check first 50 paragraphs
        text = para.text.strip()
        if not text:
            continue
            
        clean_text = remove_emojis(text)
        clean_text = re.sub(r'\s+', ' ', clean_text).strip()
        
        # Skip section headings (text ending with colon) - but allow if it's short and looks like a title
        if clean_text.endswith(':') and len(clean_text) > 100:
            continue
        
        # Skip very short text (likely headings) - but allow if it contains market name
        text_lower = clean_text.lower()
        has_market_keywords = any(kw in text_lower for kw in filename_keywords if len(kw) > 3)
        if len(clean_text) < 20 and not has_market_keywords:
            continue
        
        # Priority 1: Look for "The Global [Topic] Market" or "Global [Topic] Market" at the start
        # This is usually the actual title
        title_match = re.search(r'(?:^the\s+)?global\s+([^.]*?)\s+market', clean_text, re.I)
        if title_match:
            # Extract the title up to "Market"
            full_title_match = re.search(r'(?:^the\s+)?global\s+[^.]*?\s+market', clean_text, re.I)
            if full_title_match:
                extracted_title = full_title_match.group(0).strip()
                # Check if it contains the filename keywords (normalize both)
                text_normalized = clean_text.lower().replace('-', ' ').replace('_', ' ')
                # Check if at least some keywords match
                matching_keywords = sum(1 for kw in filename_keywords if kw in text_normalized)
                if matching_keywords >= min_keywords_needed and not _is_section_heading_title(extracted_title):
                    return _ensure_filename_start_and_year(extracted_title, filename)
        
        # Priority 1.5: Look for full title with segmentation pattern
        # Pattern: "[Market Name] Market By [Category] (...); By [Category] (...); ... Forecast, 2024–2030"
        if 'market' in text_lower and re.search(r'by\s+(?:application|product\s+type|type|end[-\s]*user|region|geography)', text_lower, re.I):
            # Check if it contains market name and forecast
            text_normalized = clean_text.lower().replace('-', ' ').replace('_', ' ')
            matching_keywords = sum(1 for kw in filename_keywords if kw in text_normalized)
            if matching_keywords >= min_keywords_needed and ('forecast' in text_lower or re.search(r'20\d{2}', clean_text)):
                # Extract title from start to forecast year or end
                year_match = re.search(r'(20\d{2}.*?20\d{2})', clean_text)
                if year_match:
                    end_pos = year_match.end()
                    full_title = clean_text[:end_pos].strip()
                    # Remove "The Global" prefix
                    full_title = re.sub(r'^(?:The\s+)?Global\s+', '', full_title, flags=re.I).strip()
                    if not _is_section_heading_title(full_title):
                        return _clean_final_title(_norm(full_title))
        
        # Priority 2: Look for patterns like "[Topic] Market" that contains filename
        text_normalized = clean_text.lower().replace('-', ' ').replace('_', ' ')
        matching_keywords = sum(1 for kw in filename_keywords if kw in text_normalized)
        
        # More flexible matching - check if text starts with capital (or quote/capital)
        starts_with_capital = clean_text and (clean_text[0].isupper() or (len(clean_text) > 1 and clean_text[0] in ['"', "'"] and clean_text[1].isupper()))
        
        if (starts_with_capital and 
            matching_keywords >= min_keywords_needed and
            'market' in text_lower and 
            len(clean_text) < 400):  # Increased length limit
            # Check if it's not a heading (doesn't end with colon if long, doesn't start with numbers)
            is_heading = re.match(r'^\d+[\.\)]\s*', clean_text)
            if not is_heading:
                # Extract up to "Market" or first sentence ending with forecast/year
                # First try to find complete title with forecast
                forecast_match = re.search(r'^(.*?market.*?forecast.*?20\d{2}.*?20\d{2})', clean_text, re.I | re.DOTALL)
                if forecast_match:
                    title_text = forecast_match.group(1).strip()
                    title_text = re.sub(r'^(?:The\s+)?Global\s+', '', title_text, flags=re.I).strip()
                    if not _is_section_heading_title(title_text):
                        return _ensure_filename_start_and_year(title_text, filename)
                
                # Otherwise extract up to "Market" or first sentence
                title_match = re.search(r'^([^.]*?market)', clean_text, re.I)
                if title_match:
                    title_text = title_match.group(1).strip()
                    # Only use if it has substantial content and not a section heading
                    if len(title_text.split()) >= 3 and not _is_section_heading_title(title_text):
                        return _ensure_filename_start_and_year(title_text, filename)
        
        # Priority 3: Look for "Forecast, 2024–2030" pattern with market name
        if re.search(r'forecast\s*[,:]\s*20\d{2}[\s\-–]20\d{2}', text_lower):
            text_normalized = clean_text.lower().replace('-', ' ').replace('_', ' ')
            matching_keywords = sum(1 for kw in filename_keywords if kw in text_normalized)
            if matching_keywords >= min_keywords_needed or 'market' in text_lower:
                # Extract title portion before forecast
                forecast_pos = re.search(r'forecast\s*[,:]', text_lower, re.I)
                if forecast_pos:
                    title_part = clean_text[:forecast_pos.end()].strip()
                    # Try to extract complete title
                    if 'market' in title_part.lower() and not _is_section_heading_title(clean_text):
                        return _ensure_filename_start_and_year(clean_text, filename)
    
        # Priority 4: Look for market name with "Market" keyword (flexible positioning)
        if matching_keywords >= min_keywords_needed and 'market' in text_lower and len(clean_text) < 200:
            # Check if it looks like a title (starts with capital, reasonable length)
            if starts_with_capital and not re.match(r'^\d+[\.\)]\s*', clean_text):
                # Extract title pattern
                # Look for: Market Name + Market [+ optional segmentation] [+ Forecast/Year]
                title_pattern = re.search(r'^((?:The\s+)?(?:Global\s+)?[^.]*?market(?:\s+by\s+[^,;.]*?)*?)(?:\s*[,;]\s*.*?forecast.*?20\d{2}.*?20\d{2})?', clean_text, re.I)
                if title_pattern:
                    title_text = title_pattern.group(1).strip()
                    title_text = re.sub(r'^(?:The\s+)?Global\s+', '', title_text, flags=re.I).strip()
                    if len(title_text.split()) >= 3 and not _is_section_heading_title(title_text):
                        return _ensure_filename_start_and_year(title_text, filename)

    # Last resort: If we found segmentation patterns but no full title, construct basic title
    # Check if we have any segmentation patterns in document
    has_segmentation = False
    for para in doc.paragraphs[:100]:
        text = para.text.strip().lower()
        if re.search(r'by\s+(?:application|product\s+type|type|end[-\s]*user|region|geography|segment)', text, re.I):
            has_segmentation = True
            break
    
    if has_segmentation:
        # Construct basic title with filename + Market + Forecast
        base_title = filename.replace('-', ' ').replace('_', ' ')
        if not base_title.lower().endswith('market'):
            base_title = f"{base_title} Market"
        return _ensure_filename_start_and_year(base_title, filename)

    return "Title Not Available"

# ------------------- Extract Description -------------------
def extract_description(docx_path):
    doc = Document(docx_path)
    html_output = []
    capture, inside_list = False, None
    last_heading = None
    used_headings = set()
    inside_regional_section = False
    inside_recent_developments_section = False
    inside_introduction_section = False
    inside_segmentation_section = False
    inside_segmentation_subheading = False
    inside_market_trends_section = False
    inside_competitive_intelligence_section = False
    inside_end_user_dynamics_section = False
    inside_regional_landscape_section = False
    inside_market_segmentation_section = False
    previous_bold_word = None  # Track previous paragraph's bold first word

    

    target_headings = [
        "introduction and strategic context",
        "market segmentation and forecast scope",
        "market trends and innovation landscape",
        "competitive intelligence and benchmarking",
        "regional landscape and adoption outlook",
        "end-user dynamics and use case",
        "recent developments + opportunities & restraints",
        "opportunities & restraints"
        
    ]
    
    # Regional headings that should only be h2 when inside "Regional Landscape" section
    regional_headings = [
        "north america",
        "north-america",
        "europe", 
        "asia pacific",
        "asia-pacific",
        "latin america",
        "latin-america",
        "latin america and middle east & africa (lamea)",
        "latin america & middle east & africa (lamea)",
        "latin america & middle east and africa (lamea)",
        "latin-america and middle east & africa (lamea)",
        "latin-america & middle east & africa (lamea)",
        "latin-america & middle east and africa (lamea)",
        "latin america and middle-east & africa (lamea)",
        "latin america & middle-east & africa (lamea)",
        "latin america & middle-east and africa (lamea)",
        "latin-america and middle-east & africa (lamea)",
        "latin-america & middle-east & africa (lamea)",
        "latin-america & middle-east and africa (lamea)",
        "latin america and middle east & africa (LAMEA)",
        "latin america & middle east & africa (LAMEA)",
        "latin america & middle east and africa (LAMEA)",
        "middle east and africa",
        "middle east & africa",
        "middle east and africa (mea)",
        "middle east & africa (mea)",
        "middle-east and africa",
        "middle-east & africa",
        "middle-east and africa (mea)",
        "middle-east & africa (mea)"
    ]
    
    # Opportunities heading that should only be h2 when inside "Recent Developments" section
    opportunities_heading = ["opportunities","restraints","Opportunities & Restraints"]
    
    # Segmentation headings now use pattern-based detection (any "By X" format)
    # No keyword list needed - automatically detects all "By X" patterns

    def clean_heading(text):
        text = remove_emojis(text.strip())
        text = re.sub(r'^[^\w]+', '', text)
        text = re.sub(r'(?i)section\s*\d+[:\-]?\s*', '', text)
        text = re.sub(r'^\d+[\.\-\)]\s*', '', text)
        text = re.sub(r'\s+', ' ', text)
        return text.lower().strip()

    def add_nbsp_safely(content=None):
        """Add &nbsp; only if the last item is not already &nbsp; and not before first heading"""
        if not html_output or html_output[-1] != "&nbsp;":
            # Check if this would be before the first heading
            if not any("<h2><strong>" in item for item in html_output):
                return  # Don't add &nbsp; before first heading
            
            # If content is provided, only add &nbsp; for long content (>= 200 chars)
            # Exception: Always add &nbsp; for numbered points (1., 2., 3., etc.)
            if content is not None and len(content) < 200:
                # Check if this is a numbered point
                if re.match(r'^\d+\.\s*', content.strip()):
                    html_output.append("&nbsp;")
                    return
                return
                
            html_output.append("&nbsp;")

    for block in doc.element.body:
        if isinstance(block, CT_P):  
            para = Paragraph(block, doc)
            text = remove_emojis(para.text.strip())
            if not text:
                continue

            cleaned = clean_heading(text)

            # Start capture
            if not capture and any(h in cleaned for h in target_headings):
                capture = True  

            # End capture - enhanced conditions
            if capture and any(end_phrase in cleaned for end_phrase in [
                "report summary, faqs, and seo schema",
                "report title",
                "report coverage table",
                "7.1. report coverage table",
                "report coverage",
                "faqs and seo schema"
            ]):
                break  

            if capture:
                content = runs_to_html(para.runs)
                matched_heading = next((h for h in target_headings if h in cleaned), None)
                
                # Check for regional headings with more flexible matching
                regional_heading = None
                for h in regional_headings:
                    if h in cleaned:
                        regional_heading = h
                        break
                
                # Special check for LAMEA - simple direct matching
                if not regional_heading and inside_regional_section:
                    # Direct check for LAMEA text pattern
                    if ("latin america" in cleaned and "middle east" in cleaned and "africa" in cleaned and 
                        ("lamea" in cleaned or "LAMEA" in text)):
                        regional_heading = "latin america and middle east & africa (lamea)"
                        print(f"DEBUG LAMEA: Direct match found for LAMEA heading")
                    
                    # Also check in content with HTML tags
                    elif ("latin america" in content.lower() and "middle east" in content.lower() and 
                          "africa" in content.lower() and ("lamea" in content.lower() or "LAMEA" in content)):
                        regional_heading = "latin america and middle east & africa (lamea)"
                        print(f"DEBUG LAMEA: Content match found for LAMEA heading")
                
                # Check for opportunities heading
                opportunities_match = next((h for h in opportunities_heading if h in cleaned), None)
                
                # Check for segmentation headings with pattern-based detection
                segmentation_heading = None
                
                # Pattern-based detection for "By X" format headings
                # ONLY check in segmentation section (between "Market Segmentation And Forecast Scope" and "Market Trends And Innovation Landscape")
                if inside_segmentation_section:
                    # Check both cleaned text and content for "By X" pattern
                    text_to_check = cleaned
                    content_to_check = content.lower()
                    
                    # Check for "by X" pattern in cleaned text
                    if re.match(r'^by\s+', text_to_check):
                        match = re.match(r'^by\s+(.+)', text_to_check)
                        if match:
                            segmentation_heading = f"by {match.group(1)}"
                            print(f"DEBUG: Found segmentation heading by pattern: '{text}' -> '{segmentation_heading}'")
                    
                    # Also check in content for bold "By X" patterns
                    elif re.search(r'<b>\s*by\s+', content_to_check):
                        bold_match = re.search(r'<b>\s*by\s+([^<]+)</b>', content_to_check)
                        if bold_match:
                            segmentation_heading = f"by {bold_match.group(1).strip()}"
                            print(f"DEBUG: Found bold segmentation heading: '{text}' -> '{segmentation_heading}'")
                
                # No fallback needed - pattern-based detection handles all "By X" formats

                if matched_heading and matched_heading not in used_headings:
                    last_heading = matched_heading
                    if matched_heading == "report coverage table":
                        last_heading = "report coverage table"  # flag set
                        continue  # ❌ skip this heading completely
                    
                    # Set flag when entering different sections
                    if matched_heading == "introduction and strategic context":
                        inside_introduction_section = True
                        inside_regional_section = False
                        inside_recent_developments_section = False
                        inside_segmentation_section = False
                    elif matched_heading == "market segmentation and forecast scope":
                        # Stop introduction spacing when next heading comes
                        inside_introduction_section = False
                        inside_regional_section = False
                        inside_recent_developments_section = False
                        inside_segmentation_section = True
                        inside_segmentation_subheading = False
                        inside_market_trends_section = False
                        inside_market_segmentation_section = True  # Start Market Segmentation logic
                    elif matched_heading == "market trends and innovation landscape":
                        inside_introduction_section = False
                        inside_regional_section = False
                        inside_recent_developments_section = False
                        inside_segmentation_section = False
                        inside_segmentation_subheading = False
                        inside_market_segmentation_section = False  # Stop Market Segmentation logic here
                        inside_market_trends_section = True
                    elif matched_heading == "competitive intelligence and benchmarking":
                        inside_introduction_section = False
                        inside_regional_section = False
                        inside_recent_developments_section = False
                        inside_segmentation_section = False
                        inside_segmentation_subheading = False
                        inside_market_trends_section = False  # Stop Market Trends logic here
                        inside_competitive_intelligence_section = True  # Start Competitive Intelligence logic
                    elif matched_heading == "regional landscape and adoption outlook":
                        inside_regional_section = True
                        inside_recent_developments_section = False
                        inside_introduction_section = False
                        inside_segmentation_section = False
                        inside_segmentation_subheading = False
                        inside_market_trends_section = False
                        inside_competitive_intelligence_section = False  # Stop Competitive Intelligence logic here
                        inside_regional_landscape_section = True  # Start Regional Landscape logic
                    elif matched_heading == "recent developments + opportunities & restraints":
                        inside_recent_developments_section = True
                        inside_regional_section = False
                        inside_introduction_section = False
                        inside_segmentation_section = False
                        inside_segmentation_subheading = False
                        inside_market_trends_section = False
                        inside_competitive_intelligence_section = False
                        inside_end_user_dynamics_section = False  # Stop End-User Dynamics logic here
                    elif matched_heading in ["end-user dynamics and use case"]:
                        inside_regional_section = False
                        inside_recent_developments_section = False
                        inside_introduction_section = False
                        inside_segmentation_section = False
                        inside_segmentation_subheading = False
                        inside_market_trends_section = False
                        inside_competitive_intelligence_section = False
                        inside_regional_landscape_section = False  # Stop Regional Landscape logic here
                        inside_end_user_dynamics_section = True  # Start End-User Dynamics logic

                    if inside_list:
                        html_output.append(f"</{inside_list}>")
                        inside_list = None

                    # ✅ Add &nbsp; before all main headings EXCEPT "Introduction And Strategic Context"
                    if matched_heading != "introduction and strategic context":
                        add_nbsp_safely()
                    
                    html_output.append(f"<h2><strong>{matched_heading.title()}</strong></h2>")
                    used_headings.add(matched_heading)
                
                # Handle regional headings only when inside regional section
                elif regional_heading and inside_regional_section and regional_heading not in used_headings:
                    # Check if it's a standalone heading (no text before or after in the same paragraph)
                    is_standalone = (
                        len(text.strip()) <= len(regional_heading) + 10 and  # Allow more extra characters for longer headings
                        text.strip().lower().startswith(regional_heading.lower()) and
                        not any(char in text for char in [',', '.', ';', ':', '!', '?'])  # No punctuation
                    )
                    
                    # Check for numbered/bulleted regional headings like "1. North America", "• Europe", etc.
                    if not is_standalone:
                        # Remove numbering and bullets to check the core heading
                        cleaned_for_regional = re.sub(r'^\d+[\.\)]\s*', '', text.strip())  # Remove "1. ", "2) ", etc.
                        cleaned_for_regional = re.sub(r'^[•\-–]\s*', '', cleaned_for_regional)  # Remove bullets
                        cleaned_for_regional = cleaned_for_regional.lower()
                        
                        # Check if the cleaned text matches any regional heading
                        for h in regional_headings:
                            if h in cleaned_for_regional or cleaned_for_regional.startswith(h):
                                # Check if it's reasonably short (not a long paragraph)
                                if len(text.strip()) <= 50 and not any(char in text for char in [';', '!', '?']):
                                    is_standalone = True
                                    print(f"DEBUG: Found numbered/bulleted regional heading: '{text}' -> '{h}'")
                                    break
                    
                    # Simple check for LAMEA standalone heading
                    if not is_standalone and regional_heading and "latin america" in regional_heading:
                        # For LAMEA, be more flexible - check if it's just the heading text
                        text_length = len(text.strip())
                        # LAMEA heading is around 50 characters, allow up to 80
                        if (text_length <= 80 and 
                            not any(char in text for char in ['.', ';', '!', '?']) and  # Allow commas and colons
                            ("latin america" in text.lower() or "latin america" in content.lower())):
                            is_standalone = True
                            print(f"DEBUG LAMEA: Simple standalone check passed for: '{text}'")
                    
                    # Additional fallback for longer regional headings
                    if not is_standalone and regional_heading and len(regional_heading) > 30:
                        # More flexible check for long regional headings
                        cleaned_text = text.strip().lower()
                        content_text = content.lower()
                        
                        # Check both original text and content (with HTML)
                        if ((regional_heading in cleaned_text and 
                             len(cleaned_text) <= len(regional_heading) + 15) or
                            ("latin america" in content_text and "middle east" in content_text and 
                             "africa" in content_text and ("lamea" in content_text or "LAMEA" in content_text))):
                            # Additional check: no sentence punctuation
                            if not any(char in text for char in [',', '.', ';', ':', '!', '?']):
                                is_standalone = True
                    
                    if is_standalone:
                        if inside_list:
                            html_output.append(f"</{inside_list}>")
                            inside_list = None

                        # ✅ Add &nbsp; before <h2>, but not after
                        add_nbsp_safely()
                        
                        # Special logic for Regional Landscape section: Add &nbsp; above regional headings
                        if inside_regional_landscape_section:
                            # Add &nbsp; above regional headings in Regional Landscape section
                            if not html_output or html_output[-1] != "&nbsp;":
                                html_output.append("&nbsp;")
                        
                        # Use original text for title case conversion instead of regional_heading
                        title_text = text.strip()
                        if "<b>" in content and "</b>" in content:
                            # Extract text from bold tags for cleaner title
                            bold_match = re.search(r'<b>(.*?)</b>', content)
                            if bold_match:
                                title_text = bold_match.group(1)
                        
                        # Clean up numbered/bulleted text for title
                        title_text = re.sub(r'^\d+[\.\)]\s*', '', title_text)  # Remove "1. ", "2) "
                        title_text = re.sub(r'^[•\-–]\s*', '', title_text)  # Remove bullets
                        title_text = title_text.strip()
                        
                        html_output.append(f"<h2><strong>{title_text}</strong></h2>")
                        used_headings.add(regional_heading)
                    else:
                        # It's part of a larger sentence, treat as normal paragraph
                        if inside_list:
                            html_output.append(f"</{inside_list}>")
                            inside_list = None
                        
                        # Special logic for Regional Landscape section: Add &nbsp; above paragraphs with regional heading + next word not bold
                        if inside_regional_landscape_section and not is_list_item(para):
                            # Check if this is a regional heading but next word is not bold
                            # Look for pattern: <b>RegionalHeading</b> followed by non-bold text
                            regional_heading_bold_match = re.match(r'^<b>([^<]+)</b>', content.strip())
                            if regional_heading_bold_match:
                                # Check if the next word after bold is not bold
                                remaining_content = content.strip()[len(regional_heading_bold_match.group(0)):].strip()
                                # If there's content after bold word, check if it doesn't start with <b>
                                if remaining_content and not remaining_content.startswith('<b>'):
                                    # Add &nbsp; above paragraphs with regional heading + next word not bold
                                    if not html_output or html_output[-1] != "&nbsp;":
                                        html_output.append("&nbsp;")
                        
                        html_output.append(f"<p style='line-height:1.6'>{content}</p>")
                
                # Handle opportunities heading only when inside recent developments section
                elif opportunities_match and inside_recent_developments_section and opportunities_match not in used_headings:
                    if inside_list:
                        html_output.append(f"</{inside_list}>")
                        inside_list = None

                    # ✅ Add &nbsp; before <h2>, but not after
                    add_nbsp_safely()
                    html_output.append(f"<h2><strong>{opportunities_match.title()}</strong></h2>")
                    used_headings.add(opportunities_match)
                
                # Handle segmentation headings (both keyword-based and pattern-based)
                elif segmentation_heading and segmentation_heading not in used_headings:
                    # Check if it's a standalone heading (no text before or after in the same paragraph)
                    is_standalone = (
                        len(text.strip()) <= len(segmentation_heading) + 10 and  # Allow more extra characters for punctuation
                        text.strip().lower().startswith(segmentation_heading.lower()) and
                        not any(char in text for char in [',', '.', ';', '!', '?'])  # No punctuation (allow colon)
                    )
                    
                    # Special case: if text ends with colon and is close to heading length, treat as standalone
                    if not is_standalone and text.strip().endswith(':') and len(text.strip()) <= len(segmentation_heading) + 5:
                        is_standalone = True
                    
                    # Special case: if text is close to heading length and contains the heading pattern, treat as standalone
                    if not is_standalone and len(text.strip()) <= len(segmentation_heading) + 5:
                        # Check if the text contains the heading pattern (for hyphen/underscore variations)
                        base_pattern = segmentation_heading.replace(' ', '').replace('-', '').replace('_', '')
                        if base_pattern in text.strip().lower().replace(' ', '').replace('-', '').replace('_', ''):
                            is_standalone = True
                    
                    # Pattern-based segmentation headings (like "By psychiatric condition")
                    if not is_standalone and segmentation_heading.startswith("by "):
                        # More flexible check for "By X" pattern headings
                        if (len(text.strip()) <= 50 and  # Reasonable length
                            ("by " in text.lower() or "<b>by " in content.lower()) and
                            not any(char in text for char in [';', '!', '?'])):  # No sentence punctuation
                            is_standalone = True
                            print(f"DEBUG: Pattern-based segmentation heading detected: '{text}'")
                    
                    
                    if is_standalone:
                        if inside_list:
                            html_output.append(f"</{inside_list}>")
                            inside_list = None

                        # Set flag that we're inside a segmentation subheading
                        inside_segmentation_subheading = True

                        # ✅ Add &nbsp; before <h2>, but not after
                        add_nbsp_safely()
                        
                        # Special logic for Market Segmentation section: Add &nbsp; above segmentation headings
                        if inside_market_segmentation_section:
                            # Add &nbsp; above segmentation headings in Market Segmentation section
                            if not html_output or html_output[-1] != "&nbsp;":
                                html_output.append("&nbsp;")
                        
                        html_output.append(f"<h2><strong>{text.strip()}</strong></h2>")
                        used_headings.add(segmentation_heading)
                    else:
                        # It's part of a larger sentence, treat as normal paragraph
                        if inside_list:
                            html_output.append(f"</{inside_list}>")
                            inside_list = None
                        html_output.append(f"<p style='line-height:1.6'>{content}</p>")

                # Subheading detection → h3
                elif re.match(r'^\d+(\.\d+)+', text.strip()):  
                    if inside_list:
                        html_output.append(f"</{inside_list}>")
                        inside_list = None
                    html_output.append(f"<h3><strong>{content}</strong></h3>")

                elif is_list_item(para):
                    if inside_list != "ul":
                        if inside_list:
                            html_output.append(f"</{inside_list}>")
                        # Don't add &nbsp; before starting a list
                        html_output.append("<ul>")
                        inside_list = "ul"

                    # ✅ CKEditor-friendly list items with p tags
                    html_output.append(f"<li><p>{content}</p></li>")

                else:
                    if inside_list:
                        html_output.append(f"</{inside_list}>")
                        inside_list = None
                    
                    # Check if this paragraph comes immediately after a main heading
                    is_after_main_heading = False
                    if html_output and "<h2><strong>" in html_output[-1]:
                        is_after_main_heading = True
                    
                    # Check if this paragraph is entirely bold (pure bold paragraph)
                    is_entirely_bold = content.strip().startswith('<b>') and content.strip().endswith('</b>')
                    
                    # Check if this paragraph has any bold content
                    has_bold_content = "<b>" in content and "</b>" in content
                    
                    # Special logic for Market Trends section: Add &nbsp; above paragraphs
                    if inside_market_trends_section and not is_list_item(para) and not is_after_main_heading:
                        should_add_nbsp = False
                        
                        # First check: If entire paragraph is bold (p tag with bold content)
                        if is_entirely_bold and len(content) < 80 and not content.strip().endswith(':'):
                            # Extract bold word from entirely bold paragraph for consecutive check
                            bold_match = re.match(r'^<b>([^<]+)</b>', content.strip())
                            if bold_match:
                                # Normalize bold word by removing common prefixes (numbers, bullets, etc.) and make lowercase
                                raw_bold_word = bold_match.group(1)
                                normalized_bold_word = re.sub(r'^\d+\.\s*', '', raw_bold_word).strip().lower()
                                current_bold_word = normalized_bold_word
                                # Don't add &nbsp; if this is the same bold word as previous paragraph
                                if current_bold_word != previous_bold_word:
                                    should_add_nbsp = True
                                # Update previous bold word for next iteration
                                previous_bold_word = current_bold_word
                        # Second check: If first word is bold (only if first check didn't match)
                        elif content.strip():
                            # Check if content starts with <b> followed by a word
                            bold_match = re.match(r'^<b>([^<]+)</b>', content.strip())
                            if bold_match:
                                # Normalize bold word by removing common prefixes (numbers, bullets, etc.) and make lowercase
                                raw_bold_word = bold_match.group(1)
                                normalized_bold_word = re.sub(r'^\d+\.\s*', '', raw_bold_word).strip().lower()
                                current_bold_word = normalized_bold_word
                                # Check if the next word after bold is not a colon
                                remaining_content = content.strip()[len(bold_match.group(0)):].strip()
                                # If there's content after bold word, check if it doesn't start with colon
                                if not remaining_content or not remaining_content.startswith(':'):
                                    # Check if this is a company name continuation (same company, different format)
                                    is_same_company = False
                                    if previous_bold_word:
                                        # Check if current word is contained in previous word or vice versa
                                        if (current_bold_word in previous_bold_word or 
                                            previous_bold_word in current_bold_word):
                                            is_same_company = True
                                            print(f"DEBUG: Same company detected: '{current_bold_word}' vs '{previous_bold_word}'")
                                    
                                    # Don't add &nbsp; if this is the same bold word or same company as previous paragraph
                                    if current_bold_word != previous_bold_word and not is_same_company:
                                        should_add_nbsp = True
                                # Update previous bold word for next iteration
                                previous_bold_word = current_bold_word
                        
                        # Special check for numbered points like "7. Company Name"
                        if not should_add_nbsp and re.match(r'^<b>\d+\.\s*', content.strip()):
                            should_add_nbsp = True
                        
                        if should_add_nbsp:
                            # Add &nbsp; above paragraphs in Market Trends section
                            if not html_output or html_output[-1] != "&nbsp;":
                                html_output.append("&nbsp;")
                    
                    # Special logic for Competitive Intelligence section: Add &nbsp; above paragraphs
                    if inside_competitive_intelligence_section and not is_list_item(para) and not is_after_main_heading:
                        should_add_nbsp = False
                        print(f"DEBUG: Competitive Intelligence section processing: {content.strip()[:30]}...")
                        
                        # Special check for numbered points like "7. Company Name"
                        if re.match(r'^<b>\d+\.\s*', content.strip()):
                            should_add_nbsp = True
                        
                        # First check: If entire paragraph is bold (p tag with bold content)
                        if is_entirely_bold and len(content) < 80 and not content.strip().endswith(':'):
                            # Extract bold word from entirely bold paragraph for consecutive check
                            bold_match = re.match(r'^<b>([^<]+)</b>', content.strip())
                            if bold_match:
                                # Normalize bold word by removing common prefixes (numbers, bullets, etc.) and make lowercase
                                raw_bold_word = bold_match.group(1)
                                normalized_bold_word = re.sub(r'^\d+\.\s*', '', raw_bold_word).strip().lower()
                                current_bold_word = normalized_bold_word
                                # Don't add &nbsp; if this is the same bold word as previous paragraph
                                if current_bold_word != previous_bold_word:
                                    should_add_nbsp = True
                                # Update previous bold word for next iteration
                                previous_bold_word = current_bold_word
                        # Second check: If first word is bold (only if first check didn't match)
                        elif content.strip():
                            # Check if content starts with <b> followed by a word
                            bold_match = re.match(r'^<b>([^<]+)</b>', content.strip())
                            if bold_match:
                                # Normalize bold word by removing common prefixes (numbers, bullets, etc.) and make lowercase
                                raw_bold_word = bold_match.group(1)
                                normalized_bold_word = re.sub(r'^\d+\.\s*', '', raw_bold_word).strip().lower()
                                current_bold_word = normalized_bold_word
                                # Check if the next word after bold is not a colon
                                remaining_content = content.strip()[len(bold_match.group(0)):].strip()
                                # If there's content after bold word, check if it doesn't start with colon
                                if not remaining_content or not remaining_content.startswith(':'):
                                    # Check if this is a company name continuation (same company, different format)
                                    is_same_company = False
                                    if previous_bold_word:
                                        # Check if current word is contained in previous word or vice versa
                                        if (current_bold_word in previous_bold_word or 
                                            previous_bold_word in current_bold_word):
                                            is_same_company = True
                                            print(f"DEBUG: Same company detected: '{current_bold_word}' vs '{previous_bold_word}'")
                                    
                                    # Don't add &nbsp; if this is the same bold word or same company as previous paragraph
                                    if current_bold_word != previous_bold_word and not is_same_company:
                                        should_add_nbsp = True
                                # Update previous bold word for next iteration
                                previous_bold_word = current_bold_word
                        
                        # Special check for numbered points like "7. Company Name"
                        if not should_add_nbsp and re.match(r'^<b>\d+\.\s*', content.strip()):
                            should_add_nbsp = True
                            print(f"DEBUG: Competitive Intelligence - Special check triggered for: {content.strip()[:50]}...")
                        
                        if should_add_nbsp:
                            # Add &nbsp; above paragraphs in Competitive Intelligence section
                            if not html_output or html_output[-1] != "&nbsp;":
                                html_output.append("&nbsp;")
                                print(f"DEBUG: Competitive Intelligence - Added &nbsp; for: {content.strip()[:30]}...")
                    
                    # Special logic for End-User Dynamics section: Add &nbsp; above paragraphs
                    if inside_end_user_dynamics_section and not is_list_item(para) and not is_after_main_heading:
                        should_add_nbsp = False
                        
                        # First check: If entire paragraph is bold (p tag with bold content)
                        if is_entirely_bold and len(content) < 80 and not content.strip().endswith(':'):
                            # Extract bold word from entirely bold paragraph for consecutive check
                            bold_match = re.match(r'^<b>([^<]+)</b>', content.strip())
                            if bold_match:
                                # Normalize bold word by removing common prefixes (numbers, bullets, etc.) and make lowercase
                                raw_bold_word = bold_match.group(1)
                                normalized_bold_word = re.sub(r'^\d+\.\s*', '', raw_bold_word).strip().lower()
                                current_bold_word = normalized_bold_word
                                # Don't add &nbsp; if this is the same bold word as previous paragraph
                                if current_bold_word != previous_bold_word:
                                    should_add_nbsp = True
                                # Update previous bold word for next iteration
                                previous_bold_word = current_bold_word
                        # Second check: If first word is bold (only if first check didn't match)
                        elif content.strip():
                            # Check if content starts with <b> followed by a word
                            bold_match = re.match(r'^<b>([^<]+)</b>', content.strip())
                            if bold_match:
                                # Normalize bold word by removing common prefixes (numbers, bullets, etc.) and make lowercase
                                raw_bold_word = bold_match.group(1)
                                normalized_bold_word = re.sub(r'^\d+\.\s*', '', raw_bold_word).strip().lower()
                                current_bold_word = normalized_bold_word
                                # Check if the next word after bold is not a colon
                                remaining_content = content.strip()[len(bold_match.group(0)):].strip()
                                # If there's content after bold word, check if it doesn't start with colon
                                if not remaining_content or not remaining_content.startswith(':'):
                                    # Check if this is a company name continuation (same company, different format)
                                    is_same_company = False
                                    if previous_bold_word:
                                        # Check if current word is contained in previous word or vice versa
                                        if (current_bold_word in previous_bold_word or 
                                            previous_bold_word in current_bold_word):
                                            is_same_company = True
                                            print(f"DEBUG: Same company detected: '{current_bold_word}' vs '{previous_bold_word}'")
                                    
                                    # Don't add &nbsp; if this is the same bold word or same company as previous paragraph
                                    if current_bold_word != previous_bold_word and not is_same_company:
                                        should_add_nbsp = True
                                # Update previous bold word for next iteration
                                previous_bold_word = current_bold_word
                        
                        # Special check for numbered points like "7. Company Name"
                        if not should_add_nbsp and re.match(r'^<b>\d+\.\s*', content.strip()):
                            should_add_nbsp = True
                            print(f"DEBUG: End-User Dynamics - Special check triggered for: {content.strip()[:50]}...")
                        
                        if should_add_nbsp:
                            # Add &nbsp; above paragraphs in End-User Dynamics section
                            if not html_output or html_output[-1] != "&nbsp;":
                                html_output.append("&nbsp;")
                                print(f"DEBUG: End-User Dynamics - Added &nbsp; for: {content.strip()[:30]}...")
                    
                    # Special logic for Regional Landscape section: Add &nbsp; above paragraphs
                    if inside_regional_landscape_section and not is_list_item(para) and not is_after_main_heading:
                        should_add_nbsp = False
                        
                        # First check: If entire paragraph is bold (p tag with bold content)
                        if is_entirely_bold and len(content) < 80 and not content.strip().endswith(':'):
                            # Extract bold word from entirely bold paragraph for consecutive check
                            bold_match = re.match(r'^<b>([^<]+)</b>', content.strip())
                            if bold_match:
                                # Normalize bold word by removing common prefixes (numbers, bullets, etc.) and make lowercase
                                raw_bold_word = bold_match.group(1)
                                normalized_bold_word = re.sub(r'^\d+\.\s*', '', raw_bold_word).strip().lower()
                                current_bold_word = normalized_bold_word
                                # Don't add &nbsp; if this is the same bold word as previous paragraph
                                if current_bold_word != previous_bold_word:
                                    should_add_nbsp = True
                                # Update previous bold word for next iteration
                                previous_bold_word = current_bold_word
                        # Second check: If first word is bold (only if first check didn't match)
                        elif content.strip():
                            # Check if content starts with <b> followed by a word
                            bold_match = re.match(r'^<b>([^<]+)</b>', content.strip())
                            if bold_match:
                                # Normalize bold word by removing common prefixes (numbers, bullets, etc.) and make lowercase
                                raw_bold_word = bold_match.group(1)
                                normalized_bold_word = re.sub(r'^\d+\.\s*', '', raw_bold_word).strip().lower()
                                current_bold_word = normalized_bold_word
                                # Check if the next word after bold is not a colon
                                remaining_content = content.strip()[len(bold_match.group(0)):].strip()
                                # If there's content after bold word, check if it doesn't start with colon
                                if not remaining_content or not remaining_content.startswith(':'):
                                    # Check if this is a company name continuation (same company, different format)
                                    is_same_company = False
                                    if previous_bold_word:
                                        # Check if current word is contained in previous word or vice versa
                                        if (current_bold_word in previous_bold_word or 
                                            previous_bold_word in current_bold_word):
                                            is_same_company = True
                                            print(f"DEBUG: Same company detected: '{current_bold_word}' vs '{previous_bold_word}'")
                                    
                                    # Don't add &nbsp; if this is the same bold word or same company as previous paragraph
                                    if current_bold_word != previous_bold_word and not is_same_company:
                                        should_add_nbsp = True
                                # Update previous bold word for next iteration
                                previous_bold_word = current_bold_word
                        
                        # Special check for numbered points like "7. Company Name"
                        if not should_add_nbsp and re.match(r'^<b>\d+\.\s*', content.strip()):
                            should_add_nbsp = True
                            print(f"DEBUG: Regional Landscape - Special check triggered for: {content.strip()[:50]}...")
                        
                        if should_add_nbsp:
                            # Add &nbsp; above paragraphs in Regional Landscape section
                            if not html_output or html_output[-1] != "&nbsp;":
                                html_output.append("&nbsp;")
                                print(f"DEBUG: Regional Landscape - Added &nbsp; for: {content.strip()[:30]}...")
                    
                    # Special logic for Market Segmentation section: Add &nbsp; above bold paragraphs
                    if inside_market_segmentation_section and is_entirely_bold and not is_after_main_heading:
                        # Apply nbsp only for content with length < 100 characters and not ending with colon
                        if (len(content) < 80 and not content.strip().endswith(':')):  # Only add if content is less than 100 characters and doesn't end with colon
                            # Add &nbsp; directly without length check for Market Segmentation bold paragraphs
                            if not html_output or html_output[-1] != "&nbsp;":
                                html_output.append("&nbsp;")
                    
                    
                    
                    # No &nbsp; before paragraphs - only after paragraphs with length >= 200
                    
                    # Add &nbsp; BEFORE the paragraph if we're in Segmentation section (only for main heading paragraphs, not sub-heading paragraphs)
                    # Check if this paragraph comes after a segmentation subheading
                    is_after_subheading = False
                    if inside_segmentation_section:
                        # Look at recent headings to see if we're after a segmentation subheading
                        for recent_item in html_output[-10:]:  # Check last 10 items
                            if "<h2><strong>By " in recent_item:
                                is_after_subheading = True
                                break
                    
                    # No &nbsp; before paragraphs - only after paragraphs with length >= 200
                    
                    html_output.append(f"<p style='line-height:1.6'>{content}</p>")
                    
                    # Add &nbsp; after paragraph only if line length >= 200 characters
                    # BUT skip this default logic when in Market Trends, Competitive Intelligence, End-User Dynamics, Regional Landscape, or Market Segmentation sections
                    if (len(content) >= 200 and not inside_market_trends_section and 
                        not inside_competitive_intelligence_section and not inside_end_user_dynamics_section and
                        not inside_regional_landscape_section and not inside_market_segmentation_section):
                        html_output.append("&nbsp;")
                    
                    # Reset subheading flag AFTER processing the paragraph
                    if inside_segmentation_subheading:
                        inside_segmentation_subheading = False
                    
        elif isinstance(block, CT_Tbl):  
            # ❌ Skip table if last heading was "report coverage table"
            if last_heading == "report coverage table":
                continue

            table = Table(block, doc)
            # Use the enhanced table styling function for consistency
            table_html = extract_table_with_style(table)
            html_output.append(table_html)

    if inside_list:
        html_output.append(f"</{inside_list}>")

    return "\n".join(html_output)

# ------------------- Helper Functions -------------------
def runs_to_html(runs):
    """Convert Word runs (bold/italic) to inline HTML."""
    parts = []
    for run in runs:
        txt = remove_emojis(run.text.strip())
        if not txt:
            continue
        if run.bold and run.italic:
            parts.append(f"<b><i>{txt}</i></b>")
        elif run.bold:
            parts.append(f"<b>{txt}</b>")
        elif run.italic:
            parts.append(f"<i>{txt}</i>")
        else:
            parts.append(txt)
    return " ".join(parts).strip()

 # ------------------- FAQ Schema + Methodology -------------------
def _get_text(docx_path):
    doc = Document(docx_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())

def _extract_json_block(text, type_name):
    pat = re.compile(r'"@type"\s*:\s*"' + re.escape(type_name) + r'"')
    m = pat.search(text)
    if not m:
        return ""
    start_idx = text.rfind("{", 0, m.start())
    if start_idx == -1:
        return ""
    depth, i, n = 0, start_idx, len(text)
    block_chars = []
    while i < n:
        ch = text[i]
        block_chars.append(ch)
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                break
        i += 1
    return "".join(block_chars).strip()

def extract_faq_schema(docx_path):
    text = _get_text(docx_path)
    return _extract_json_block(text, "FAQPage")

def extract_methodology_from_faqschema(docx_path):
    faq_schema_str = extract_faq_schema(docx_path)  
    if not faq_schema_str:
        return ""   
    try:
        # Clean the JSON string by removing extra whitespace and newlines
        cleaned_json = re.sub(r'\s+', ' ', faq_schema_str.strip())
        faq_data = json.loads(cleaned_json)
    except json.JSONDecodeError:
        return ""   
    faqs = []
    q_count = 0
    for item in faq_data.get("mainEntity", []):
        q_count += 1
        question = item.get("name", "").strip()
        answer = item.get("acceptedAnswer", {}).get("text", "").strip()
        if question and answer:
            faqs.append(
                f"<p><strong>Q{q_count}: {html.escape(question)}</strong><br>"
                f"A{q_count}: {html.escape(answer)}</p>"
            )
    return "\n".join(faqs)

# ------------------- Report Coverage -------------------
def extract_report_coverage_table_with_style(docx_path):
    doc = Document(docx_path)
    print(f"DEBUG: Found {len(doc.tables)} tables in document")  # Debug log
    
    for table_idx, table in enumerate(doc.tables):
        if len(table.rows) == 0:
            continue
            
        first_row_text = " ".join([c.text.strip().lower() for c in table.rows[0].cells])
        print(f"DEBUG: Table {table_idx} first row: {first_row_text}")  # Debug log
        
        # Check if this looks like a report coverage table
        is_report_table = (
            "report attribute" in first_row_text or 
            "report coverage table" in first_row_text or
            "forecast period" in first_row_text or
            "market size" in first_row_text or
            "revenue forecast" in first_row_text or
            ("forecast" in first_row_text and "period" in first_row_text) or
            ("market" in first_row_text and "size" in first_row_text)
        )
        
        if is_report_table:
            print(f"DEBUG: Found report coverage table at index {table_idx}")  # Debug log
            html_parts = []
            html_parts.append('<h2><strong>7.1. Report Coverage Table</strong></h2>')
            html_parts.append('')
            html_parts.append('<table cellspacing=0 style=\'border-collapse:collapse; width:100%\'>')
            html_parts.append('        <tbody>')
            
            for r_idx, row in enumerate(table.rows):
                html_parts.append('            <tr>')
                
                # Process each cell in the row
                for c_idx, cell in enumerate(row.cells):
                    text = remove_emojis(cell.text.strip())
                    
                    # Determine cell styling based on position
                    if r_idx == 0:  # Header row
                        if c_idx == 0:  # First column
                            cell_style = "background-color:#4472c4; border-bottom:1px solid #4472c4; border-left:1px solid #4472c4; border-right:none; border-top:1px solid #4472c4; vertical-align:top; width:195px"
                        else:  # Second column
                            cell_style = "background-color:#4472c4; border-bottom:1px solid #4472c4; border-left:none; border-right:1px solid #4472c4; border-top:1px solid #4472c4; vertical-align:top; width:370px"
                        
                        html_parts.append(f'                <td style=\'{cell_style}\'>')
                        html_parts.append(f'                <p><strong>{text}</strong></p>')
                        html_parts.append(f'                </td>')
                    
                    else:  # Data rows
                        # Alternate row colors
                        bg_color = "#d9e2f3" if r_idx % 2 == 1 else ""
                        
                        if c_idx == 0:  # First column
                            if bg_color:
                                cell_style = f"background-color:{bg_color}; border-bottom:1px solid #8eaadb; border-left:1px solid #8eaadb; border-right:1px solid #8eaadb; border-top:none; vertical-align:top; width:195px"
                            else:
                                cell_style = "border-bottom:1px solid #8eaadb; border-left:1px solid #8eaadb; border-right:1px solid #8eaadb; border-top:none; vertical-align:top; width:195px"
                        else:  # Second column
                            if bg_color:
                                cell_style = f"background-color:{bg_color}; border-bottom:1px solid #8eaadb; border-left:none; border-right:1px solid #8eaadb; border-top:none; vertical-align:top; width:370px"
                            else:
                                cell_style = "border-bottom:1px solid #8eaadb; border-left:none; border-right:1px solid #8eaadb; border-top:none; vertical-align:top; width:370px"
                        
                        html_parts.append(f'                <td style=\'{cell_style}\'>')
                        
                        # Both columns are bold
                        html_parts.append(f'                <p><strong>{text}</strong></p>')
                        
                        html_parts.append(f'                </td>')
                
                html_parts.append('            </tr>')
            
            html_parts.append('        </tbody>')
            html_parts.append('</table>')
            print(f"DEBUG: Generated HTML for report coverage table")  # Debug log
            return "\n".join(html_parts)
    
    print("DEBUG: No report coverage table found")  # Debug log
    return ""

# ------------------- Extra Extractors -------------------
def extract_meta_description(docx_path):
    doc = Document(docx_path)
    capture = False
    for para in doc.paragraphs:
        text = para.text.strip()
        low = text.lower()
        if not capture and ("introduction" in low):
            capture = True
            continue
        if capture and text:
            return text
    return ""

def extract_seo_title(docx_path):
    doc = Document(docx_path)
    file_name = os.path.splitext(os.path.basename(docx_path))[0]
    revenue_forecast = ""

    def normalize_label(txt: str) -> str:
        t = txt.strip().lower()
        t = re.sub(r"\s+", " ", t)
        t = t.replace("forecast by", "forecast in")
        t = t.replace("forecasts in", "forecast in")
        t = t.replace("forecast (", "forecast in ")
        t = t.replace(")", "")
        return t

    for table in doc.tables:
        if not table.rows or not table.rows[0].cells:
            continue
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if "report attribute" in headers and "details" in headers:
            attr_idx = headers.index("report attribute")
            details_idx = headers.index("details")
            for row in table.rows[1:]:
                attr_raw = row.cells[attr_idx].text.strip()
                attr = normalize_label(attr_raw)
                details = row.cells[details_idx].text.strip()
                # Match any variant that implies revenue/market size forecast for 2030
                attr_lower = attr_raw.lower()
                if (("revenue forecast" in attr and (" 2030" in attr or "forecast in" in attr)) or \
                   ("revenue forecast" in attr_lower and "2030" in attr_raw)) or \
                   (("market size forecast" in attr_lower or "market size" in attr_lower) and "2030" in attr_raw):
                    revenue_forecast = re.sub(r"USD", "$", details, flags=re.I).strip()
                    break
        if revenue_forecast:
            break

    # SEO title: market name with spaces (no underscores)
    title_name = file_name.replace("_", " ")
    if revenue_forecast:
        return f"{title_name} Size ({revenue_forecast}) 2030"
    return title_name

def extract_breadcrumb_text(docx_path):
    file_name = os.path.splitext(os.path.basename(docx_path))[0]
    revenue_forecast = ""
    doc = Document(docx_path)

    def normalize_label(txt: str) -> str:
        t = txt.strip().lower()
        t = re.sub(r"\s+", " ", t)
        t = t.replace("forecast by", "forecast in")
        t = t.replace("forecasts in", "forecast in")
        t = t.replace("forecast (", "forecast in ")
        t = t.replace(")", "")
        return t

    for table in doc.tables:
        if not table.rows or not table.rows[0].cells:
            continue
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if "report attribute" in headers and "details" in headers:
            attr_idx = headers.index("report attribute")
            details_idx = headers.index("details")
            for row in table.rows[1:]:
                attr_raw = row.cells[attr_idx].text.strip()
                attr = normalize_label(attr_raw)
                details = row.cells[details_idx].text.strip()
                # Match any variant that implies revenue/market size forecast for 2030
                attr_lower = attr_raw.lower()
                if (("revenue forecast" in attr and (" 2030" in attr or "forecast in" in attr)) or \
                   ("revenue forecast" in attr_lower and "2030" in attr_raw)) or \
                   (("market size forecast" in attr_lower or "market size" in attr_lower) and "2030" in attr_raw):
                    revenue_forecast = re.sub(r"USD", "$", details, flags=re.I).strip()
                    break
        if revenue_forecast:
            break

    # Breadcrumb: market name with spaces (no underscores)
    title_name = file_name.replace("_", " ")
    if revenue_forecast:
        return f"{title_name} Report 2030"
    return title_name

def extract_sku_code(docx_path):
    import re
    sku_code = os.path.splitext(os.path.basename(docx_path))[0]
    
    # Apply new SKU processing rules:
    # 1. Replace "and" with space (case insensitive)
    processed_sku = re.sub(r'\band\b', ' ', sku_code, flags=re.IGNORECASE)
    
    # 2. Remove "Global" word (case insensitive)
    processed_sku = re.sub(r'\bglobal\b', '', processed_sku, flags=re.IGNORECASE)
    
    # 3. Remove parentheses and content inside, replace with space
    processed_sku = re.sub(r'\([^)]*\)', ' ', processed_sku)
    
    # 4. Replace "- and" with single space (case insensitive)
    processed_sku = re.sub(r'\s*-\s*and\b', ' ', processed_sku, flags=re.IGNORECASE)
    
    # 5. Replace hyphens with space
    processed_sku = processed_sku.replace('-', ' ')
    
    # 6. Remove all special characters except letters, numbers and spaces
    processed_sku = re.sub(r'[^a-zA-Z0-9\s]', ' ', processed_sku)
    
    # 7. Clean up multiple spaces and trim
    processed_sku = re.sub(r'\s+', ' ', processed_sku).strip()
    
    # 8. Convert to lowercase
    processed_sku = processed_sku.lower()
    
    return processed_sku

def extract_sku_url(docx_path):
    import re
    sku_code = os.path.splitext(os.path.basename(docx_path))[0]
    
    # Apply same SKU processing rules as extract_sku_code:
    # 1. Replace & with space
    processed_sku = sku_code.replace('&', ' ')
    
    # 2. Replace - with space  
    processed_sku = processed_sku.replace('-', ' ')
    
    # 3. Replace "and" with space (case insensitive)
    processed_sku = re.sub(r'\band\b', ' ', processed_sku, flags=re.IGNORECASE)
    
    # 4. Remove parentheses and content inside, replace with space
    processed_sku = re.sub(r'\([^)]*\)', ' ', processed_sku)
    
    # 5. Clean up multiple spaces and trim
    processed_sku = re.sub(r'\s+', ' ', processed_sku).strip()
    
    # 6. Convert to lowercase
    processed_sku = processed_sku.lower()
    
    return processed_sku

def extract_breadcrumb_schema(docx_path):
    text = _get_text(docx_path)
    breadcrumb_json = _extract_json_block(text, "BreadcrumbList")
    
    if not breadcrumb_json:
        return ""
    
    # Get SKU code
    sku_code = extract_sku_code(docx_path)
    
    # Add "market" to SKU code if it doesn't end with "market"
    if not sku_code.lower().endswith("market"):
        sku_code = sku_code + " market"
    
    # Replace spaces with hyphens for URL
    sku_code = sku_code.replace(" ", "-")
    
    # Use regex to directly replace the URL in position 3
    import re
    
    # Pattern to match the item URL in position 3
    pattern = r'("item":\s*"https://www\.strategicmarketresearch\.com/market-report/)[^"]*(")'
    
    # Use a lambda function to avoid group reference issues
    def replace_url(match):
        return match.group(1) + sku_code + match.group(2)
    
    # Replace the URL
    modified_json = re.sub(pattern, replace_url, breadcrumb_json)
    
    return modified_json

# ------------------- Merge -------------------
def merge_description_and_coverage(docx_path):
    try:
        desc_html = extract_description(docx_path) or ""
        coverage_html = extract_report_coverage_table_with_style(docx_path) or ""
        merged_html = desc_html + "\n\n" + coverage_html if (desc_html or coverage_html) else ""
        return merged_html
    except Exception as e:
        return f"ERROR: {e}"      
    return [text[i:i+limit] for i in range(0, len(text), limit)]      
